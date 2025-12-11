# utils/document_converter/document_converter_utils.py - DOCUMENT CONVERTER UTILITIES
# Dedicated utilities for universal document conversion
import json
import logging
import os
import shutil
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# Universal Security Framework Integration
try:
    from security.core.sanitizers import sanitize_filename
    from security.core.validators import validate_content, validate_filename

    SECURITY_FRAMEWORK_AVAILABLE = True
except ImportError:
    SECURITY_FRAMEWORK_AVAILABLE = False

# Document conversion dependencies
try:
    import pypandoc

    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False

try:
    import docx.enum.style
    import docx.shared
    from docx import Document as DocxDocument
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import weasyprint

    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A3, A4, legal, letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from bs4 import BeautifulSoup

    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import pdfminer.six
    import pdfplumber

    PDF_PROCESSING_AVAILABLE = True
except ImportError:
    PDF_PROCESSING_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image as PILImage

    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from odf import style
    from odf.opendocument import OpenDocumentText
    from odf.text import H, P

    ODF_AVAILABLE = True
except ImportError:
    ODF_AVAILABLE = False

try:
    import markdown
    import markdown.extensions.tables
    import markdown.extensions.toc

    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    import ebooklib
    from ebooklib import epub

    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False

try:
    import docx2txt

    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

try:
    import win32com.client

    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False

try:
    import chardet

    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False

# Initialize logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentConverter:
    """
    Universal document converter supporting:
    - Document format detection and validation
    - Multi-format conversions (DOCX, PDF, HTML, Markdown, RTF, ODT, TXT, EPUB)
    - Batch processing and file merging
    - Advanced options (OCR, formatting preservation, etc.)
    """

    def __init__(self):
        self.temp_dirs = []  # Track temporary directories for cleanup

        # Check for required dependencies
        self.dependencies = self._check_dependencies()

        # Supported format mappings
        self.input_formats = {
            "docx",
            "doc",
            "rtf",
            "odt",
            "txt",
            "md",
            "html",
            "htm",
            "epub",
            "pdf",
        }

        self.output_formats = {"pdf", "docx", "html", "markdown", "rtf", "odt", "txt"}

    def __del__(self):
        """Cleanup temporary directories on object destruction"""
        self.cleanup_temp_dirs()

    def cleanup_temp_dirs(self):
        """Clean up all temporary directories created by this instance"""
        for temp_dir in self.temp_dirs:
            if os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                except Exception as e:
                    logger.warning(f"Failed to cleanup temp dir {temp_dir}: {e}")
        self.temp_dirs.clear()

    def _check_dependencies(self) -> Dict[str, bool]:
        """Check availability of all dependencies"""
        return {
            "pandoc": PANDOC_AVAILABLE,
            "python_docx": PYTHON_DOCX_AVAILABLE,
            "weasyprint": WEASYPRINT_AVAILABLE,
            "reportlab": REPORTLAB_AVAILABLE,
            "bs4": BS4_AVAILABLE,
            "pdf_processing": PDF_PROCESSING_AVAILABLE,
            "ocr": OCR_AVAILABLE,
            "odf": ODF_AVAILABLE,
            "markdown": MARKDOWN_AVAILABLE,
            "epub": EPUB_AVAILABLE,
            "docx2txt": DOCX2TXT_AVAILABLE,
            "win32com": WIN32COM_AVAILABLE,
        }

    def is_pdf_conversion_available(self) -> bool:
        """Check if PDF conversion is available"""
        return (
            self.dependencies["weasyprint"]
            or self.dependencies["reportlab"]
            or self.dependencies["pandoc"]
        )

    def is_docx_conversion_available(self) -> bool:
        """Check if DOCX conversion is available"""
        return self.dependencies["python_docx"] or self.dependencies["pandoc"]

    def is_html_conversion_available(self) -> bool:
        """Check if HTML conversion is available"""
        return self.dependencies["bs4"] or self.dependencies["pandoc"]

    def is_markdown_conversion_available(self) -> bool:
        """Check if markdown conversion is available"""
        return self.dependencies["markdown"]

    def _read_file_with_encoding_detection(self, file_path: str) -> Tuple[str, str]:
        """
        Read a file with robust encoding detection.

        Returns:
            Tuple[str, str]: (file_content, detected_encoding)
        """
        # First, try to detect encoding using chardet if available
        if CHARDET_AVAILABLE:
            try:
                with open(file_path, "rb") as f:
                    raw_data = f.read()

                # Use chardet for encoding detection
                detection_result = chardet.detect(raw_data)
                detected_encoding = detection_result.get("encoding", "utf-8")
                confidence = detection_result.get("confidence", 0.0)

                # If confidence is high, try the detected encoding
                if confidence > 0.7 and detected_encoding:
                    try:
                        content = raw_data.decode(detected_encoding)
                        return content, detected_encoding
                    except (UnicodeDecodeError, LookupError):
                        pass

            except Exception as e:
                logger.warning(f"Encoding detection failed for {file_path}: {e}")
        else:
            logger.debug(
                f"chardet not available, using fallback encoding detection for {file_path}"
            )

        # Fallback: try common encodings in order of preference
        common_encodings = [
            "utf-8",
            "utf-8-sig",  # UTF-8 with BOM
            "windows-1252",  # Common Windows encoding
            "cp1252",  # Another name for Windows-1252
            "iso-8859-1",  # Latin-1
            "ascii",
            "utf-16",
            "utf-16le",
            "utf-16be",
        ]

        for encoding in common_encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                logger.info(f"Successfully read {file_path} with encoding: {encoding}")
                return content, encoding
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                logger.warning(f"Error reading {file_path} with {encoding}: {e}")
                continue

        # Last resort: read as binary and decode with error handling
        try:
            # Read as binary if we haven't already done so
            if not CHARDET_AVAILABLE or "raw_data" not in locals():
                with open(file_path, "rb") as f:
                    raw_data = f.read()

            # Try to decode with UTF-8 and replace/ignore errors
            content = raw_data.decode("utf-8", errors="replace")
            logger.warning(f"Used UTF-8 with error replacement for {file_path}")
            return content, "utf-8-replaced"

        except Exception as e:
            # Ultimate fallback
            try:
                content = raw_data.decode("latin-1", errors="ignore")
                logger.error(f"Used Latin-1 with error ignoring for {file_path}: {e}")
                return content, "latin-1-ignored"
            except Exception as final_error:
                logger.error(f"Failed to read file {file_path}: {final_error}")
                raise UnicodeDecodeError(
                    "unknown",
                    b"",
                    0,
                    1,
                    f"Could not decode file {file_path} with any encoding",
                )

    def validate_file_security(self, file_path: str) -> Tuple[bool, List[str]]:
        """Validate file using universal security framework"""
        if not SECURITY_FRAMEWORK_AVAILABLE:
            logger.warning(
                "Security framework not available - performing basic validation"
            )
            return self._basic_file_validation(file_path)

        try:
            # Read file content for security validation
            with open(file_path, "rb") as f:
                file_content = f.read()

            # Determine file type
            file_ext = Path(file_path).suffix.lower().lstrip(".")

            # Perform security validation
            is_safe, issues = validate_content(file_content, file_ext)

            if not is_safe:
                logger.error(f"Security validation failed for {file_path}: {issues}")
                return False, issues

            logger.info(f"Security validation passed for {file_path}")
            return True, []

        except Exception as e:
            logger.error(f"Security validation error for {file_path}: {e}")
            return False, [f"Security validation error: {str(e)}"]

    def _basic_file_validation(self, file_path: str) -> Tuple[bool, List[str]]:
        """Basic file validation when security framework is not available"""
        try:
            if not os.path.exists(file_path):
                return False, ["File does not exist"]

            file_size = os.path.getsize(file_path)
            if file_size == 0:
                return False, ["File is empty"]

            if file_size > 100 * 1024 * 1024:  # 100MB limit
                return False, ["File too large"]

            # Basic file type validation based on extension
            file_ext = Path(file_path).suffix.lower().lstrip(".")
            if file_ext not in self.input_formats:
                return False, [f"Unsupported file format: {file_ext}"]

            return True, []

        except Exception as e:
            return False, [f"Basic validation error: {str(e)}"]

    def detect_document_format(self, file_path: str) -> str:
        """Detect document format from file"""
        try:
            file_ext = Path(file_path).suffix.lower().lstrip(".")

            # Extension-based detection
            format_map = {
                "docx": "docx",
                "doc": "doc",
                "rtf": "rtf",
                "odt": "odt",
                "txt": "txt",
                "md": "markdown",
                "html": "html",
                "htm": "html",
                "epub": "epub",
                "pdf": "pdf",
            }

            detected_format = format_map.get(file_ext, "unknown")

            # Content-based validation for specific formats
            if detected_format == "unknown":
                detected_format = self._detect_format_by_content(file_path)

            return detected_format

        except Exception as e:
            logger.error(f"Error detecting document format for {file_path}: {e}")
            return "unknown"

    def _detect_format_by_content(self, file_path: str) -> str:
        """Detect format by examining file content"""
        try:
            with open(file_path, "rb") as f:
                header = f.read(1024)

            # PDF signature
            if header.startswith(b"%PDF-"):
                return "pdf"

            # ZIP-based formats (DOCX, ODT, EPUB)
            if header.startswith(b"PK\x03\x04"):
                # Could be DOCX, ODT, or EPUB
                try:
                    import zipfile

                    with zipfile.ZipFile(file_path, "r") as zf:
                        file_list = zf.namelist()
                        if "word/document.xml" in file_list:
                            return "docx"
                        elif "content.xml" in file_list:
                            return "odt"
                        elif "META-INF/container.xml" in file_list:
                            return "epub"
                except:
                    pass

            # RTF signature
            if header.startswith(b"{\\rtf"):
                return "rtf"

            # HTML content
            if b"<html" in header.lower() or b"<!doctype html" in header.lower():
                return "html"

            # Default to text
            return "txt"

        except Exception as e:
            logger.error(f"Error detecting format by content: {e}")
            return "unknown"

    def convert_single_document(
        self,
        input_path: str,
        output_path: str,
        output_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert a single document to the specified format"""
        try:
            # Validate input file
            is_safe, security_issues = self.validate_file_security(input_path)
            if not is_safe:
                return (
                    False,
                    f"Security validation failed: {', '.join(security_issues)}",
                )

            # Detect input format
            input_format = self.detect_document_format(input_path)
            if input_format == "unknown":
                return False, "Unable to detect input document format"

            logger.info(f"Converting {input_format} to {output_format}: {input_path}")

            # Special handling for DOC files - convert to DOCX first for better format preservation
            if input_format == "doc" and output_format != "txt":
                return self._convert_doc_via_docx(
                    input_path, output_path, output_format, options
                )

            # Route to appropriate conversion method
            if output_format == "pdf":
                return self._convert_to_pdf(
                    input_path, output_path, input_format, options
                )
            elif output_format == "docx":
                return self._convert_to_docx(
                    input_path, output_path, input_format, options
                )
            elif output_format == "html":
                return self._convert_to_html(
                    input_path, output_path, input_format, options
                )
            elif output_format == "markdown":
                return self._convert_to_markdown(
                    input_path, output_path, input_format, options
                )
            elif output_format == "rtf":
                return self._convert_to_rtf(
                    input_path, output_path, input_format, options
                )
            elif output_format == "odt":
                return self._convert_to_odt(
                    input_path, output_path, input_format, options
                )
            elif output_format == "txt":
                return self._convert_to_txt(
                    input_path, output_path, input_format, options
                )
            else:
                return False, f"Unsupported output format: {output_format}"

        except Exception as e:
            logger.error(f"Error converting document: {e}")
            return False, f"Conversion error: {str(e)}"

    def merge_and_convert_documents(
        self,
        input_paths: List[str],
        output_path: str,
        output_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Merge multiple documents and convert to specified format"""
        try:
            if len(input_paths) < 2:
                return False, "At least 2 files required for merging"

            # Create temporary merged file
            temp_dir = tempfile.mkdtemp(prefix="document_merge_")
            self.temp_dirs.append(temp_dir)

            # Merge strategy depends on target format and input types
            if output_format in ["txt", "markdown"]:
                # Text-based merging
                merged_path = os.path.join(temp_dir, "merged.txt")
                success, error = self._merge_as_text(input_paths, merged_path, options)
            elif output_format == "html":
                # HTML merging
                merged_path = os.path.join(temp_dir, "merged.html")
                success, error = self._merge_as_html(input_paths, merged_path, options)
            else:
                # Convert all to intermediate format first, then merge
                merged_path = os.path.join(temp_dir, "merged.html")
                success, error = self._merge_via_html(input_paths, merged_path, options)

            if not success:
                return False, error

            # Convert merged document to final format
            if output_format == "txt" and merged_path.endswith(".txt"):
                # Already in target format
                shutil.copy2(merged_path, output_path)
                return True, ""
            elif output_format == "html" and merged_path.endswith(".html"):
                # Already in target format
                shutil.copy2(merged_path, output_path)
                return True, ""
            else:
                # Convert merged file to target format
                return self.convert_single_document(
                    merged_path, output_path, output_format, options
                )

        except Exception as e:
            logger.error(f"Error merging documents: {e}")
            return False, f"Merge error: {str(e)}"

    def _merge_as_text(
        self, input_paths: List[str], output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Merge documents as plain text"""
        try:
            merged_content = []
            separator = "\n" + "=" * 80 + "\n"

            for input_path in input_paths:
                # Convert each document to text first
                temp_txt = tempfile.NamedTemporaryFile(
                    mode="w+", suffix=".txt", delete=False
                )
                temp_txt.close()

                # Extract text and add to merged content
                success, error = self._convert_to_txt(
                    input_path,
                    temp_txt.name,
                    self.detect_document_format(input_path),
                    options,
                )
                if success:
                    try:
                        content, _ = self._read_file_with_encoding_detection(
                            temp_txt.name
                        )
                        merged_content.append(f"Source: {Path(input_path).name}")
                        merged_content.append(separator)
                        merged_content.append(content)
                        merged_content.append(separator)
                    except Exception as e:
                        logger.error(f"Error reading merged text content: {e}")
                        merged_content.append(f"Source: {Path(input_path).name}")
                        merged_content.append(separator)
                        merged_content.append(f"Error reading content: {e}")
                        merged_content.append(separator)

                os.unlink(temp_txt.name)

            if not merged_content:
                return False, "No content could be extracted for merging"

            # Write merged content
            with open(
                output_path, "w", encoding=options.get("text_encoding", "utf-8")
            ) as f:
                f.write("\n".join(merged_content))

            return True, ""

        except Exception as e:
            logger.error(f"Error merging as text: {e}")
            return False, str(e)

    def _merge_as_html(
        self, input_paths: List[str], output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Merge documents as HTML"""
        try:
            html_parts = []
            html_parts.append("<!DOCTYPE html>")
            html_parts.append(
                '<html><head><meta charset="utf-8"><title>Merged Document</title>'
            )
            html_parts.append(
                "<style>body { font-family: Arial, sans-serif; margin: 40px; }</style>"
            )
            html_parts.append("</head><body>")

            for i, input_path in enumerate(input_paths):
                # Convert each document to HTML first
                temp_html = tempfile.NamedTemporaryFile(
                    mode="w+", suffix=".html", delete=False
                )
                temp_html.close()

                success, error = self._convert_to_html(
                    input_path,
                    temp_html.name,
                    self.detect_document_format(input_path),
                    options,
                )
                if success:
                    try:
                        content, encoding_used = (
                            self._read_file_with_encoding_detection(temp_html.name)
                        )
                        logger.debug(
                            f"Read HTML content with encoding: {encoding_used}"
                        )
                    except Exception as e:
                        logger.error(f"Failed to read HTML file {temp_html.name}: {e}")
                        content = f"<p>Error reading content: {e}</p>"

                    # Extract body content if it's a full HTML document
                    if "<body>" in content:
                        body_start = content.find("<body>") + 6
                        body_end = content.find("</body>")
                        if body_end > body_start:
                            content = content[body_start:body_end]

                    html_parts.append(
                        f"<h1>Document {i + 1}: {Path(input_path).name}</h1>"
                    )
                    html_parts.append("<hr>")
                    html_parts.append(content)
                    html_parts.append('<div style="page-break-after: always;"></div>')

                os.unlink(temp_html.name)

            html_parts.append("</body></html>")

            # Write merged HTML
            with open(output_path, "w", encoding="utf-8") as f:
                f.write("\n".join(html_parts))

            return True, ""

        except Exception as e:
            logger.error(f"Error merging as HTML: {e}")
            return False, str(e)

    def _merge_via_html(
        self, input_paths: List[str], output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Merge documents by converting all to HTML first"""
        return self._merge_as_html(input_paths, output_path, options)

    def _convert_to_pdf(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert document to PDF"""
        try:
            pdf_password = options.get("pdf_password", "").strip()
            temp_pdf_path = None

            # If password protection is needed, use a temporary file first
            if pdf_password:
                temp_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
                temp_pdf.close()
                temp_pdf_path = temp_pdf.name
                actual_output = temp_pdf_path
            else:
                actual_output = output_path

            conversion_success = False
            conversion_error = ""

            # Try methods in order of formatting preservation quality
            # For DOCX files, try COM automation first for best results
            if input_format == "docx" and self.dependencies["win32com"]:
                try:
                    success, error = self._convert_docx_to_pdf_com(
                        input_path, actual_output, options
                    )
                    if success:
                        conversion_success = True
                    else:
                        logger.warning(f"COM PDF conversion failed: {error}")
                        conversion_error = error
                except Exception as com_error:
                    logger.warning(f"COM PDF conversion error: {com_error}")
                    conversion_error = str(com_error)

            # Try Pandoc for good format support
            if (
                not conversion_success
                and self.dependencies["pandoc"]
                and input_format in ["docx", "html", "markdown", "rtf", "odt"]
            ):
                try:
                    success, error = self._convert_to_pdf_pandoc(
                        input_path, actual_output, input_format, options
                    )
                    if success:
                        conversion_success = True
                    else:
                        conversion_error = error
                except Exception as pandoc_error:
                    logger.warning(f"Pandoc PDF conversion failed: {pandoc_error}")
                    conversion_error = str(pandoc_error)

            # Try WeasyPrint (good for HTML)
            if not conversion_success and self.dependencies["weasyprint"]:
                try:
                    success, error = self._convert_to_pdf_weasyprint(
                        input_path, actual_output, input_format, options
                    )
                    if success:
                        conversion_success = True
                    else:
                        conversion_error = error
                except Exception as weasy_error:
                    logger.warning(f"WeasyPrint PDF conversion failed: {weasy_error}")
                    conversion_error = str(weasy_error)

            # Try ReportLab as fallback (already supports password)
            if not conversion_success and self.dependencies["reportlab"]:
                try:
                    success, error = self._convert_to_pdf_reportlab(
                        input_path, actual_output, input_format, options
                    )
                    if success:
                        conversion_success = True
                    else:
                        conversion_error = error
                except Exception as reportlab_error:
                    logger.warning(
                        f"ReportLab PDF conversion failed: {reportlab_error}"
                    )
                    conversion_error = str(reportlab_error)

            if not conversion_success:
                if temp_pdf_path and os.path.exists(temp_pdf_path):
                    os.unlink(temp_pdf_path)
                return False, conversion_error or "No PDF conversion library available"

            # Apply password protection if needed and not already applied by ReportLab
            if pdf_password and temp_pdf_path:
                success, error = self._apply_pdf_password_protection(
                    temp_pdf_path, output_path, pdf_password
                )
                os.unlink(temp_pdf_path)  # Clean up temp file
                if not success:
                    return False, f"Failed to apply password protection: {error}"

            return True, ""

        except Exception as e:
            return False, f"PDF conversion error: {str(e)}"

    def _apply_pdf_password_protection(
        self, input_pdf_path: str, output_pdf_path: str, password: str
    ) -> Tuple[bool, str]:
        """Apply password protection to an existing PDF file"""
        try:
            # Try using PyPDF2 first
            try:
                from PyPDF2 import PdfReader, PdfWriter

                # Read the input PDF
                reader = PdfReader(input_pdf_path)
                writer = PdfWriter()

                # Copy all pages to writer
                for page in reader.pages:
                    writer.add_page(page)

                # Add password protection
                writer.encrypt(password, password, use_128bit=True)

                # Write the protected PDF
                with open(output_pdf_path, "wb") as output_file:
                    writer.write(output_file)

                return True, ""

            except Exception as pypdf2_error:
                logger.warning(f"PyPDF2 password protection failed: {pypdf2_error}")

            # Try using reportlab to recreate the PDF with password
            if self.dependencies["reportlab"]:
                try:
                    # Extract text from the original PDF first
                    temp_txt = tempfile.NamedTemporaryFile(
                        mode="w+", suffix=".txt", delete=False
                    )
                    temp_txt.close()

                    # Use pdfplumber or fitz to extract text
                    extracted_text = ""

                    # Try pdfplumber first
                    if self.dependencies["pdf_processing"]:
                        try:
                            import pdfplumber

                            with pdfplumber.open(input_pdf_path) as pdf:
                                for page in pdf.pages:
                                    text = page.extract_text()
                                    if text:
                                        extracted_text += text + "\n"
                        except Exception:
                            pass

                    # Fallback to fitz (PyMuPDF)
                    if not extracted_text.strip():
                        try:
                            import fitz

                            pdf_document = fitz.open(input_pdf_path)
                            for page_num in range(pdf_document.page_count):
                                page = pdf_document.load_page(page_num)
                                extracted_text += page.get_text() + "\n"
                            pdf_document.close()
                        except Exception:
                            pass

                    if extracted_text.strip():
                        # Write extracted text to temp file
                        with open(temp_txt.name, "w", encoding="utf-8") as f:
                            f.write(extracted_text)

                        # Use ReportLab to create a new password-protected PDF
                        doc = SimpleDocTemplate(output_pdf_path, pagesize=A4)
                        styles = getSampleStyleSheet()
                        story = []

                        # Split content into paragraphs
                        paragraphs = extracted_text.split("\n\n")
                        for para in paragraphs:
                            if para.strip():
                                story.append(Paragraph(para.strip(), styles["Normal"]))
                                story.append(Spacer(1, 12))

                        # Build PDF first without password
                        doc.build(story)

                        # Then apply password protection using PyPDF2 if available
                        try:
                            from PyPDF2 import PdfReader, PdfWriter

                            # Read the just-created PDF
                            reader = PdfReader(output_pdf_path)
                            writer = PdfWriter()

                            # Copy all pages
                            for page in reader.pages:
                                writer.add_page(page)

                            # Add password protection
                            writer.encrypt(password, password, use_128bit=True)

                            # Overwrite with protected version
                            with open(output_pdf_path, "wb") as output_file:
                                writer.write(output_file)

                        except Exception as pypdf_error:
                            logger.warning(
                                f"PyPDF2 password protection in ReportLab fallback failed: {pypdf_error}"
                            )
                        os.unlink(temp_txt.name)
                        return True, ""

                except Exception as reportlab_error:
                    logger.warning(
                        f"ReportLab password protection failed: {reportlab_error}"
                    )

            # If all methods fail, just copy the file without password protection
            logger.warning(
                "Password protection not available, copying PDF without protection"
            )
            shutil.copy2(input_pdf_path, output_pdf_path)
            return (
                True,
                "Password protection not available, PDF saved without protection",
            )

        except Exception as e:
            logger.error(f"PDF password protection error: {e}")
            return False, str(e)

    def _convert_to_pdf_weasyprint(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert to PDF using WeasyPrint"""
        try:
            # First convert to HTML if needed
            if input_format != "html":
                temp_html = tempfile.NamedTemporaryFile(
                    mode="w+", suffix=".html", delete=False
                )
                temp_html.close()
                success, error = self._convert_to_html(
                    input_path, temp_html.name, input_format, options
                )
                if not success:
                    os.unlink(temp_html.name)
                    return False, f"HTML conversion failed: {error}"
                html_path = temp_html.name
            else:
                html_path = input_path

            # Convert HTML to PDF
            html_doc = weasyprint.HTML(filename=html_path)
            html_doc.write_pdf(output_path)

            if input_format != "html":
                os.unlink(html_path)

            return True, ""

        except Exception as e:
            logger.error(f"WeasyPrint conversion error: {e}")
            return False, str(e)

    def _convert_to_pdf_reportlab(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert to PDF using ReportLab"""
        try:
            # This is a simplified implementation
            # A full implementation would handle various input formats properly

            # Extract text content first
            temp_txt = tempfile.NamedTemporaryFile(
                mode="w+", suffix=".txt", delete=False
            )
            temp_txt.close()

            success, error = self._convert_to_txt(
                input_path, temp_txt.name, input_format, options
            )
            if not success:
                os.unlink(temp_txt.name)
                return False, f"Text extraction failed: {error}"

            # Read text content with robust encoding detection
            try:
                content, encoding_used = self._read_file_with_encoding_detection(
                    temp_txt.name
                )
                logger.debug(f"Read text content with encoding: {encoding_used}")
            except Exception as e:
                logger.error(f"Failed to read text file for PDF conversion: {e}")
                return False, f"Failed to read text content: {e}"

            os.unlink(temp_txt.name)

            # Create PDF with ReportLab
            page_size = A4
            if options.get("pdf_page_size") == "Letter":
                page_size = letter
            elif options.get("pdf_page_size") == "Legal":
                page_size = legal
            elif options.get("pdf_page_size") == "A3":
                page_size = A3

            # Create PDF document with optional password protection
            doc = SimpleDocTemplate(output_path, pagesize=page_size)
            styles = getSampleStyleSheet()
            story = []

            # Split content into paragraphs
            paragraphs = content.split("\n\n")
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), styles["Normal"]))
                    story.append(Spacer(1, 12))

            # Build PDF with optional password
            pdf_password = options.get("pdf_password", "")
            if pdf_password:
                # Build the PDF first, then apply encryption
                doc.build(story)

                # Apply password protection using PyPDF2
                success, error = self._apply_pdf_password_protection(
                    output_path, output_path, pdf_password
                )
                if not success:
                    logger.warning(f"Password protection failed: {error}")
            else:
                doc.build(story)

            return True, ""

        except Exception as e:
            logger.error(f"ReportLab conversion error: {e}")
            return False, str(e)

    def _convert_to_pdf_pandoc(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert to PDF using Pandoc"""
        try:
            # Map internal formats to pandoc formats
            format_map = {
                "docx": "docx",
                # Note: DOC format not supported by Pandoc
                "html": "html",
                "markdown": "markdown",
                "txt": "plain",
                "rtf": "rtf",
                "odt": "odt",
            }

            pandoc_input_format = format_map.get(input_format, "plain")

            # Use pandoc to convert
            extra_args = []
            if options.get("pdf_page_size"):
                page_size = options["pdf_page_size"].lower()
                # Map page sizes to correct geometry values
                size_map = {
                    "a4": "a4paper",
                    "a3": "a3paper",
                    "letter": "letterpaper",
                    "legal": "legalpaper",
                }
                geometry_size = size_map.get(page_size, "a4paper")
                extra_args.extend(["-V", f"geometry:{geometry_size}"])

            # Add PDF engine for better compatibility
            extra_args.extend(["--pdf-engine=xelatex"])

            try:
                # Try with geometry settings first
                pypandoc.convert_file(
                    input_path,
                    "pdf",
                    format=pandoc_input_format,
                    outputfile=output_path,
                    extra_args=extra_args,
                )
                return True, ""
            except Exception as geometry_error:
                if (
                    "geometry" in str(geometry_error).lower()
                    or "keyval" in str(geometry_error).lower()
                ):
                    logger.warning(
                        f"Geometry package error, trying without geometry: {geometry_error}"
                    )
                    # Try without geometry settings
                    basic_args = ["--pdf-engine=xelatex"]
                    pypandoc.convert_file(
                        input_path,
                        "pdf",
                        format=pandoc_input_format,
                        outputfile=output_path,
                        extra_args=basic_args,
                    )
                    return True, ""
                else:
                    raise geometry_error

        except Exception as e:
            logger.error(f"Pandoc PDF conversion error: {e}")
            return False, str(e)

    def _convert_to_docx(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert document to DOCX"""
        try:
            # If input is already DOCX, just copy it
            if input_format == "docx":
                shutil.copy2(input_path, output_path)
                return True, ""

            # Special handling for PDF input
            elif input_format == "pdf":
                return self._convert_pdf_to_docx(input_path, output_path, options)
            elif self.dependencies["python_docx"] and input_format in [
                "txt",
                "html",
                "markdown",
            ]:
                return self._convert_to_docx_python_docx(
                    input_path, output_path, input_format, options
                )
            elif self.dependencies["pandoc"]:
                return self._convert_to_docx_pandoc(
                    input_path, output_path, input_format, options
                )
            else:
                return False, "No DOCX conversion library available"

        except Exception as e:
            return False, f"DOCX conversion error: {str(e)}"

    def _convert_to_docx_python_docx(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert to DOCX using python-docx"""
        try:
            doc = DocxDocument()

            if input_format == "txt":
                # Plain text conversion with encoding detection
                try:
                    content, encoding_used = self._read_file_with_encoding_detection(
                        input_path
                    )
                    logger.debug(f"Read TXT file with encoding: {encoding_used}")
                except Exception as e:
                    logger.error(f"Failed to read TXT file {input_path}: {e}")
                    return False, f"Failed to read text file: {e}"

                paragraphs = content.split("\n")
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
                    else:
                        doc.add_paragraph("")  # Empty line

            elif input_format == "html":
                # HTML conversion (simplified) with encoding detection
                try:
                    html_content, encoding_used = (
                        self._read_file_with_encoding_detection(input_path)
                    )
                    logger.debug(f"Read HTML file with encoding: {encoding_used}")
                except Exception as e:
                    logger.error(f"Failed to read HTML file {input_path}: {e}")
                    return False, f"Failed to read HTML file: {e}"

                if self.dependencies["bs4"]:
                    soup = BeautifulSoup(html_content, "html.parser")

                    # Extract text content maintaining some structure
                    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
                        heading = doc.add_heading(
                            element.get_text().strip(), level=int(element.name[1])
                        )

                    for element in soup.find_all("p"):
                        doc.add_paragraph(element.get_text().strip())
                else:
                    # Fallback: strip HTML tags manually
                    import re

                    text_content = re.sub(r"<[^>]+>", "", html_content)
                    doc.add_paragraph(text_content.strip())

            elif input_format == "markdown":
                # Markdown conversion (simplified) with encoding detection
                try:
                    md_content, encoding_used = self._read_file_with_encoding_detection(
                        input_path
                    )
                    logger.debug(f"Read Markdown file with encoding: {encoding_used}")
                except Exception as e:
                    logger.error(f"Failed to read Markdown file {input_path}: {e}")
                    return False, f"Failed to read Markdown file: {e}"

                # Basic markdown parsing
                lines = md_content.split("\n")
                for line in lines:
                    line = line.strip()
                    if line.startswith("#"):
                        # Heading
                        level = len(line) - len(line.lstrip("#"))
                        text = line.lstrip("# ").strip()
                        if text:
                            doc.add_heading(text, level=min(level, 6))
                    elif line:
                        # Regular paragraph
                        doc.add_paragraph(line)
                    else:
                        # Empty line
                        doc.add_paragraph("")

            doc.save(output_path)
            return True, ""

        except Exception as e:
            logger.error(f"python-docx conversion error: {e}")
            return False, str(e)

    def _convert_pdf_to_docx(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Convert PDF to DOCX with OCR support"""
        try:
            # First extract text from PDF
            temp_txt = tempfile.NamedTemporaryFile(
                mode="w+", suffix=".txt", delete=False
            )
            temp_txt.close()

            # Extract text with OCR if requested
            success, error = self._extract_text_from_pdf(
                input_path, temp_txt.name, options
            )
            if not success:
                os.unlink(temp_txt.name)
                return False, f"Failed to extract text from PDF: {error}"

            # Now convert the extracted text to DOCX
            if self.dependencies["python_docx"]:
                try:
                    doc = DocxDocument()

                    # Read the extracted text with encoding detection
                    try:
                        content, encoding_used = (
                            self._read_file_with_encoding_detection(temp_txt.name)
                        )
                        logger.debug(
                            f"Read extracted PDF text with encoding: {encoding_used}"
                        )
                    except Exception as e:
                        logger.error(f"Failed to read extracted PDF text: {e}")
                        return False, f"Failed to read extracted text: {e}"

                    # Add document title
                    doc.add_heading("Converted from PDF", 0)

                    # Process content - try to maintain structure
                    paragraphs = content.split("\n\n")
                    for para in paragraphs:
                        if para.strip():
                            # Check if it looks like a heading (all caps or starts with number)
                            if para.isupper() or (
                                len(para) < 100 and para[0].isdigit()
                            ):
                                doc.add_heading(para.strip(), level=2)
                            else:
                                doc.add_paragraph(para.strip())

                    # Add metadata if requested
                    if options.get("include_metadata"):
                        doc.add_page_break()
                        doc.add_heading("Document Information", 1)
                        doc.add_paragraph(f"Source: {Path(input_path).name}")
                        doc.add_paragraph(
                            f"Converted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                        )
                        if options.get("ocr_scan"):
                            doc.add_paragraph(
                                "Note: This document was processed with OCR"
                            )

                    # Save the document
                    doc.save(output_path)
                    os.unlink(temp_txt.name)
                    return True, ""

                except Exception as docx_error:
                    os.unlink(temp_txt.name)
                    logger.error(f"DOCX creation error: {docx_error}")
                    return False, f"Failed to create DOCX: {str(docx_error)}"
            else:
                os.unlink(temp_txt.name)
                return (
                    False,
                    "python-docx library not available for PDF to DOCX conversion",
                )

        except Exception as e:
            logger.error(f"PDF to DOCX conversion error: {e}")
            return False, f"PDF to DOCX conversion failed: {str(e)}"

    def _convert_to_docx_pandoc(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert to DOCX using Pandoc"""
        try:
            format_map = {
                # Note: DOC format not supported by Pandoc
                "html": "html",
                "markdown": "markdown",
                "txt": "plain",
                "rtf": "rtf",
                "odt": "odt",
                # Note: PDF as input is not supported by Pandoc for DOCX output
            }

            pandoc_input_format = format_map.get(input_format, "plain")

            pypandoc.convert_file(
                input_path, "docx", format=pandoc_input_format, outputfile=output_path
            )

            return True, ""

        except Exception as e:
            logger.error(f"Pandoc DOCX conversion error: {e}")
            return False, str(e)

    def _convert_to_html(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert document to HTML with enhanced quality"""
        try:
            if input_format == "html":
                # Already HTML, enhance it if needed
                return self._enhance_existing_html(input_path, output_path, options)

            # Try methods in order of quality
            # 1. COM automation for DOCX (best quality)
            if input_format == "docx" and self.dependencies["win32com"]:
                try:
                    success, error = self._convert_docx_to_html_com(
                        input_path, output_path, options
                    )
                    if success:
                        return True, ""
                    else:
                        logger.warning(f"COM HTML conversion failed: {error}")
                except Exception as com_error:
                    logger.warning(f"COM HTML conversion error: {com_error}")

            # 2. Enhanced Pandoc conversion
            if self.dependencies["pandoc"]:
                return self._convert_to_html_pandoc_enhanced(
                    input_path, output_path, input_format, options
                )

            # 3. Basic conversion as fallback
            else:
                return self._convert_to_html_basic_enhanced(
                    input_path, output_path, input_format, options
                )

        except Exception as e:
            return False, f"HTML conversion error: {str(e)}"

    def _convert_to_html_pandoc_enhanced(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Enhanced Pandoc HTML conversion with better formatting"""
        try:
            format_map = {
                "docx": "docx",
                # Note: DOC format not supported by Pandoc
                "markdown": "markdown",
                "txt": "plain",
                "rtf": "rtf",
                "odt": "odt",
            }

            # DOC files need special handling
            if input_format == "doc":
                # Convert DOC to DOCX first
                temp_docx = tempfile.NamedTemporaryFile(
                    mode="w+b", suffix=".docx", delete=False
                )
                temp_docx.close()
                success, error = self._convert_doc_to_docx_com(
                    input_path, temp_docx.name, options
                )
                if success:
                    # Use the DOCX file for conversion
                    try:
                        pypandoc.convert_file(
                            temp_docx.name,
                            "html5",
                            format="docx",
                            outputfile=output_path,
                            extra_args=extra_args,
                        )
                        os.unlink(temp_docx.name)
                        # Post-process the HTML for additional enhancements
                        self._enhance_pandoc_html(output_path, options)
                        return True, ""
                    except Exception as e:
                        os.unlink(temp_docx.name)
                        raise e
                else:
                    # Fallback to text conversion
                    return self._convert_to_html_basic_enhanced(
                        input_path, output_path, input_format, options
                    )

            pandoc_input_format = format_map.get(input_format)
            if not pandoc_input_format:
                # Format not supported by Pandoc, use basic conversion
                return self._convert_to_html_basic_enhanced(
                    input_path, output_path, input_format, options
                )

            # Enhanced Pandoc arguments for better HTML output
            extra_args = [
                "--standalone",  # Create complete HTML document
                "--section-divs",  # Wrap sections in divs
                "--table-of-contents",  # Add TOC if headings present
                "--toc-depth=6",  # Include all heading levels
                "--number-sections",  # Number sections
                "--highlight-style=pygments",  # Syntax highlighting
                "--metadata",
                'title="Converted Document"',
            ]

            # Add CSS for enhanced styling
            if options.get("preserve_formatting", True):
                extra_args.extend(
                    ["--css=data:text/css;base64," + self._get_embedded_css_base64()]
                )

            # Convert with enhanced options
            pypandoc.convert_file(
                input_path,
                "html5",
                format=pandoc_input_format,
                outputfile=output_path,
                extra_args=extra_args,
            )

            # Post-process the HTML for additional enhancements
            self._enhance_pandoc_html(output_path, options)

            return True, ""

        except Exception as e:
            logger.error(f"Enhanced Pandoc HTML conversion error: {e}")
            return False, str(e)

    def _get_embedded_css_base64(self) -> str:
        """Get base64 encoded CSS for embedding in HTML"""
        import base64

        css_content = self._get_enhanced_html_css({})
        return base64.b64encode(css_content.encode("utf-8")).decode("ascii")

    def _enhance_pandoc_html(self, html_path: str, options: Dict[str, Any]):
        """Enhance Pandoc-generated HTML"""
        try:
            html_content, encoding_used = self._read_file_with_encoding_detection(
                html_path
            )
            logger.debug(f"Read HTML for enhancement with encoding: {encoding_used}")

            if self.dependencies["bs4"]:
                soup = BeautifulSoup(html_content, "html.parser")

                # Add meta tags for better rendering
                head = soup.find("head")
                if head:
                    # Add responsive viewport
                    if not soup.find("meta", {"name": "viewport"}):
                        viewport_meta = soup.new_tag("meta")
                        viewport_meta.attrs["name"] = "viewport"
                        viewport_meta.attrs["content"] = (
                            "width=device-width, initial-scale=1.0"
                        )
                        head.insert(0, viewport_meta)

                    # Add character encoding
                    if not soup.find("meta", {"charset": True}):
                        charset_meta = soup.new_tag("meta")
                        charset_meta.attrs["charset"] = "utf-8"
                        head.insert(0, charset_meta)

                # Enhance images with lazy loading
                for img in soup.find_all("img"):
                    img["loading"] = "lazy"
                    img["class"] = img.get("class", []) + ["enhanced-image"]

                # Enhance code blocks
                for code in soup.find_all("code"):
                    code["class"] = code.get("class", []) + ["enhanced-code"]

                # Enhance tables
                for table in soup.find_all("table"):
                    table["class"] = table.get("class", []) + ["enhanced-table"]

                # Write enhanced HTML
                with open(html_path, "w", encoding="utf-8") as f:
                    f.write(str(soup))

        except Exception as e:
            logger.warning(f"Pandoc HTML enhancement failed: {e}")

    def _convert_to_html_basic_enhanced(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Enhanced basic HTML conversion without Pandoc"""
        try:
            if input_format == "txt":
                return self._convert_txt_to_html_enhanced(
                    input_path, output_path, options
                )
            elif input_format == "markdown":
                return self._convert_markdown_to_html_enhanced(
                    input_path, output_path, options
                )
            elif input_format == "docx":
                return self._convert_docx_to_html_basic(
                    input_path, output_path, options
                )
            elif input_format == "doc":
                # Extract text from DOC and convert to HTML
                temp_txt = tempfile.NamedTemporaryFile(
                    mode="w+", suffix=".txt", delete=False
                )
                temp_txt.close()
                success, error = self._extract_text_from_doc(
                    input_path, temp_txt.name, options
                )
                if success:
                    result = self._convert_txt_to_html_enhanced(
                        temp_txt.name, output_path, options
                    )
                    os.unlink(temp_txt.name)
                    return result
                else:
                    os.unlink(temp_txt.name)
                    return False, f"Failed to extract text from DOC file: {error}"
            else:
                return (
                    False,
                    f"No enhanced basic HTML converter available for {input_format}",
                )

        except Exception as e:
            logger.error(f"Enhanced basic HTML conversion error: {e}")
            return False, str(e)

    def _convert_txt_to_html_enhanced(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Convert plain text to enhanced HTML"""
        try:
            content, encoding_used = self._read_file_with_encoding_detection(input_path)
            logger.debug(
                f"Read text file for HTML conversion with encoding: {encoding_used}"
            )
        except Exception as e:
            logger.error(f"Failed to read text file for HTML conversion: {e}")
            return False, f"Failed to read text file: {e}"

        # Process text content for better HTML presentation
        paragraphs = content.split("\n\n")
        html_body = ""

        for para in paragraphs:
            para = para.strip()
            if para:
                # Detect potential headings (lines in ALL CAPS or with specific patterns)
                lines = para.split("\n")
                if len(lines) == 1 and len(para) < 100:
                    if para.isupper() or para.endswith(":"):
                        html_body += (
                            f'<h2 class="enhanced-heading level-2">{para}</h2>\n'
                        )
                        continue

                # Convert to paragraph with line breaks preserved
                para_html = "<br>".join(lines)
                html_body += f"<p>{para_html}</p>\n"

        if not html_body.strip():
            html_body = f'<pre class="enhanced-code">{content}</pre>'

        enhanced_css = self._get_enhanced_html_css(options)
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>{enhanced_css}</style>
</head>
<body>
    <div class="document-container">
        <h1 class="enhanced-heading level-1">Converted Document</h1>
        {html_body}
    </div>
</body>
</html>"""

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        return True, ""

    def _convert_markdown_to_html_enhanced(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Convert Markdown to enhanced HTML"""
        try:
            md_content, encoding_used = self._read_file_with_encoding_detection(
                input_path
            )
            logger.debug(
                f"Read Markdown file for HTML conversion with encoding: {encoding_used}"
            )
        except Exception as e:
            logger.error(f"Failed to read Markdown file for HTML conversion: {e}")
            return False, f"Failed to read Markdown file: {e}"

        if self.dependencies["markdown"]:
            # Use full-featured markdown with extensions
            md = markdown.Markdown(
                extensions=[
                    "extra",
                    "codehilite",
                    "toc",
                    "tables",
                    "fenced_code",
                    "footnotes",
                    "smarty",
                ]
            )
            html_body = md.convert(md_content)
            toc = getattr(md, "toc", "")
        else:
            # Basic markdown conversion
            html_body = self._basic_markdown_to_html(md_content)
            toc = ""

        enhanced_css = self._get_enhanced_html_css(options)

        # Add table of contents if available
        toc_section = (
            f'<div class="table-of-contents"><h2>Table of Contents</h2>{toc}</div>'
            if toc
            else ""
        )

        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>{enhanced_css}</style>
    <style>
    .table-of-contents {{
        background-color: #f8f9fa;
        padding: 20px;
        margin-bottom: 30px;
        border-radius: 8px;
        border-left: 4px solid #3498db;
    }}
    .table-of-contents h2 {{
        margin-top: 0;
        color: #2c3e50;
    }}
    .document-container {{
        max-width: 1200px;
        margin: 0 auto;
        padding: 20px;
    }}
    </style>
</head>
<body>
    <div class="document-container">
        {toc_section}
        {html_body}
    </div>
</body>
</html>"""

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        return True, ""

    def _basic_markdown_to_html(self, md_content: str) -> str:
        """Basic markdown to HTML conversion without external libraries"""
        import re

        html = md_content

        # Headers
        html = re.sub(
            r"^# (.*)",
            r'<h1 class="enhanced-heading level-1">\1</h1>',
            html,
            flags=re.MULTILINE,
        )
        html = re.sub(
            r"^## (.*)",
            r'<h2 class="enhanced-heading level-2">\1</h2>',
            html,
            flags=re.MULTILINE,
        )
        html = re.sub(
            r"^### (.*)",
            r'<h3 class="enhanced-heading level-3">\1</h3>',
            html,
            flags=re.MULTILINE,
        )

        # Bold and italic
        html = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", html)
        html = re.sub(r"\*(.*?)\*", r"<em>\1</em>", html)

        # Code
        html = re.sub(r"`(.*?)`", r'<code class="enhanced-code">\1</code>', html)

        # Links
        html = re.sub(r"\[(.*?)\]\((.*?)\)", r'<a href="\2">\1</a>', html)

        # Paragraphs
        paragraphs = html.split("\n\n")
        html = "\n".join(
            f"<p>{p.strip()}</p>" if not p.strip().startswith("<") else p
            for p in paragraphs
            if p.strip()
        )

        return html

    def _convert_docx_to_html_basic(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Basic DOCX to HTML conversion using python-docx"""
        if not self.dependencies["python_docx"]:
            return False, "python-docx not available"

        try:
            doc = DocxDocument(input_path)

            html_body = ""

            # Extract paragraphs and basic formatting
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Try to detect heading style
                    if paragraph.style.name.startswith("Heading"):
                        level = paragraph.style.name.replace("Heading ", "")
                        try:
                            level_num = int(level)
                            html_body += f'<h{level_num} class="enhanced-heading level-{level_num}">{text}</h{level_num}>\n'
                        except:
                            html_body += (
                                f'<h2 class="enhanced-heading level-2">{text}</h2>\n'
                            )
                    else:
                        html_body += f"<p>{text}</p>\n"

            # Extract tables
            for table in doc.tables:
                html_body += '<table class="enhanced-table">\n'
                for i, row in enumerate(table.rows):
                    if i == 0:  # Header row
                        html_body += "  <thead><tr>\n"
                        for cell in row.cells:
                            html_body += f"    <th>{cell.text.strip()}</th>\n"
                        html_body += "  </tr></thead>\n  <tbody>\n"
                    else:
                        html_body += "  <tr>\n"
                        for cell in row.cells:
                            html_body += f"    <td>{cell.text.strip()}</td>\n"
                        html_body += "  </tr>\n"
                if len(table.rows) > 0:
                    html_body += "  </tbody>\n"
                html_body += "</table>\n"

            enhanced_css = self._get_enhanced_html_css(options)
            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>{enhanced_css}</style>
</head>
<body>
    <div class="document-container">
        {html_body}
    </div>
</body>
</html>"""

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(html_content)

            return True, ""

        except Exception as e:
            logger.error(f"Basic DOCX to HTML conversion error: {e}")
            return False, str(e)

    def _enhance_existing_html(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Enhance existing HTML file with better styling and structure"""
        try:
            with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                html_content = f.read()

            if self.dependencies["bs4"]:
                soup = BeautifulSoup(html_content, "html.parser")

                # Ensure proper HTML5 structure
                if not soup.html:
                    soup.wrap(soup.new_tag("html", lang="en"))

                # Enhance head section
                head = soup.find("head")
                if not head:
                    head = soup.new_tag("head")
                    soup.html.insert(0, head)

                # Add meta tags
                if not soup.find("meta", {"charset": True}):
                    charset_meta = soup.new_tag("meta")
                    charset_meta.attrs["charset"] = "utf-8"
                    head.insert(0, charset_meta)

                if not soup.find("meta", {"name": "viewport"}):
                    viewport_meta = soup.new_tag("meta")
                    viewport_meta.attrs["name"] = "viewport"
                    viewport_meta.attrs["content"] = (
                        "width=device-width, initial-scale=1.0"
                    )
                    head.append(viewport_meta)

                # Add enhanced CSS
                enhanced_css = self._get_enhanced_html_css(options)
                style_tag = soup.new_tag("style")
                style_tag.string = enhanced_css
                head.append(style_tag)

                # Enhance existing elements
                for table in soup.find_all("table"):
                    table["class"] = table.get("class", []) + ["enhanced-table"]

                for i in range(1, 7):
                    for heading in soup.find_all(f"h{i}"):
                        heading["class"] = heading.get("class", []) + [
                            f"enhanced-heading",
                            f"level-{i}",
                        ]

                for img in soup.find_all("img"):
                    img["class"] = img.get("class", []) + ["enhanced-image"]
                    img["loading"] = "lazy"

                for code in soup.find_all("code"):
                    code["class"] = code.get("class", []) + ["enhanced-code"]

                # Write enhanced HTML
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(str(soup))
            else:
                # Simple enhancement without BeautifulSoup
                enhanced_css = self._get_enhanced_html_css(options)
                if "<head>" in html_content:
                    html_content = html_content.replace(
                        "<head>", f"<head>\n<style>{enhanced_css}</style>"
                    )
                else:
                    html_content = f"<style>{enhanced_css}</style>\n" + html_content

                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(html_content)

            return True, ""

        except Exception as e:
            logger.error(f"HTML enhancement error: {e}")
            return False, str(e)

    def _convert_to_html_basic(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Basic HTML conversion without Pandoc"""
        try:
            if input_format == "txt":
                # Use robust encoding detection
                content, encoding_used = self._read_file_with_encoding_detection(
                    input_path
                )
                logger.debug(
                    f"Read text file for basic HTML conversion with encoding: {encoding_used}"
                )

                # Convert plain text to HTML
                html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Converted Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        pre {{ white-space: pre-wrap; }}
    </style>
</head>
<body>
    <pre>{content}</pre>
</body>
</html>"""

                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(html_content)

                return True, ""

            elif input_format == "markdown" and self.dependencies["markdown"]:
                # Use robust encoding detection
                md_content, encoding_used = self._read_file_with_encoding_detection(
                    input_path
                )
                logger.debug(
                    f"Read Markdown file for basic HTML conversion with encoding: {encoding_used}"
                )

                # Convert Markdown to HTML
                md = markdown.Markdown(extensions=["extra", "toc"])
                html_body = md.convert(md_content)

                html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Converted Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
    </style>
</head>
<body>
    {html_body}
</body>
</html>"""

                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(html_content)

                return True, ""

            else:
                return False, f"No basic HTML converter available for {input_format}"

        except Exception as e:
            logger.error(f"Basic HTML conversion error: {e}")
            return False, str(e)

    def _convert_to_markdown(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert document to Markdown"""
        try:
            if input_format == "markdown":
                # Already Markdown, just copy
                shutil.copy2(input_path, output_path)
                return True, ""

            if self.dependencies["pandoc"]:
                return self._convert_to_markdown_pandoc(
                    input_path, output_path, input_format, options
                )
            else:
                return self._convert_to_markdown_basic(
                    input_path, output_path, input_format, options
                )

        except Exception as e:
            return False, f"Markdown conversion error: {str(e)}"

    def _convert_to_markdown_pandoc(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert to Markdown using Pandoc"""
        try:
            format_map = {
                "docx": "docx",
                # Note: DOC format not supported by Pandoc
                "html": "html",
                "txt": "plain",
                "rtf": "rtf",
                "odt": "odt",
            }

            # DOC files need special handling
            if input_format == "doc":
                # Convert DOC to text first since Pandoc doesn't support DOC
                temp_txt = tempfile.NamedTemporaryFile(
                    mode="w+", suffix=".txt", delete=False
                )
                temp_txt.close()
                success, error = self._extract_text_from_doc(
                    input_path, temp_txt.name, {}
                )
                if success:
                    try:
                        pypandoc.convert_file(
                            temp_txt.name,
                            "markdown",
                            format="plain",
                            outputfile=output_path,
                        )
                        os.unlink(temp_txt.name)
                        return True, ""
                    except Exception as e:
                        os.unlink(temp_txt.name)
                        return False, f"Pandoc conversion failed: {str(e)}"
                else:
                    os.unlink(temp_txt.name)
                    return False, f"Failed to extract text from DOC: {error}"

            pandoc_input_format = format_map.get(input_format)
            if not pandoc_input_format:
                return False, f"Format '{input_format}' not supported by Pandoc"

            pypandoc.convert_file(
                input_path,
                "markdown",
                format=pandoc_input_format,
                outputfile=output_path,
            )

            return True, ""

        except Exception as e:
            logger.error(f"Pandoc Markdown conversion error: {e}")
            return False, str(e)

    def _convert_to_markdown_basic(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Basic Markdown conversion without Pandoc"""
        try:
            if input_format == "txt":
                # Plain text to Markdown (minimal conversion)
                # Use robust encoding detection
                content, encoding_used = self._read_file_with_encoding_detection(
                    input_path
                )
                logger.debug(
                    f"Read text file for basic Markdown conversion with encoding: {encoding_used}"
                )

                # Simple conversion: wrap in code blocks to preserve formatting
                md_content = f"# Converted Document\n\n```\n{content}\n```\n"

                with open(
                    output_path, "w", encoding=options.get("text_encoding", "utf-8")
                ) as f:
                    f.write(md_content)

                return True, ""

            elif input_format == "html" and self.dependencies["bs4"]:
                # Use robust encoding detection
                html_content, encoding_used = self._read_file_with_encoding_detection(
                    input_path
                )
                logger.debug(
                    f"Read HTML file for basic Markdown conversion with encoding: {encoding_used}"
                )

                soup = BeautifulSoup(html_content, "html.parser")

                md_lines = []

                # Convert headings
                for i in range(1, 7):
                    for heading in soup.find_all(f"h{i}"):
                        md_lines.append(
                            "#" * i + " " + heading.get_text().strip() + "\n"
                        )

                # Convert paragraphs
                for p in soup.find_all("p"):
                    md_lines.append(p.get_text().strip() + "\n\n")

                # If no structured content found, extract all text
                if not md_lines:
                    md_lines.append("# Converted Document\n\n")
                    md_lines.append(soup.get_text().strip())

                with open(
                    output_path, "w", encoding=options.get("text_encoding", "utf-8")
                ) as f:
                    f.write("".join(md_lines))

                return True, ""

            else:
                return (
                    False,
                    f"No basic Markdown converter available for {input_format}",
                )

        except Exception as e:
            logger.error(f"Basic Markdown conversion error: {e}")
            return False, str(e)

    def _convert_to_rtf(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert document to RTF"""
        try:
            if input_format == "rtf":
                # Already RTF, just copy
                shutil.copy2(input_path, output_path)
                return True, ""

            if self.dependencies["pandoc"]:
                return self._convert_to_rtf_pandoc(
                    input_path, output_path, input_format, options
                )
            else:
                return self._convert_to_rtf_basic(
                    input_path, output_path, input_format, options
                )

        except Exception as e:
            return False, f"RTF conversion error: {str(e)}"

    def _convert_to_rtf_pandoc(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert to RTF using Pandoc"""
        try:
            format_map = {
                "docx": "docx",
                # Note: DOC format not supported by Pandoc
                "html": "html",
                "markdown": "markdown",
                "txt": "plain",
                "odt": "odt",
            }

            pandoc_input_format = format_map.get(input_format, "plain")

            pypandoc.convert_file(
                input_path, "rtf", format=pandoc_input_format, outputfile=output_path
            )

            return True, ""

        except Exception as e:
            logger.error(f"Pandoc RTF conversion error: {e}")
            return False, str(e)

    def _convert_to_rtf_basic(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Basic RTF conversion without Pandoc"""
        try:
            if input_format == "txt":
                # Use robust encoding detection
                content, encoding_used = self._read_file_with_encoding_detection(
                    input_path
                )
                logger.debug(
                    f"Read text file for basic RTF conversion with encoding: {encoding_used}"
                )

                # Create basic RTF document
                rtf_content = (
                    r"""{\rtf1\ansi\deff0
{\fonttbl{\f0 Times New Roman;}}
\f0\fs24 """
                    + content.replace("\n", r"\par ")
                    + r"}"
                )

                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(rtf_content)

                return True, ""

            else:
                return False, f"No basic RTF converter available for {input_format}"

        except Exception as e:
            logger.error(f"Basic RTF conversion error: {e}")
            return False, str(e)

    def _convert_to_odt(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert document to ODT"""
        try:
            if input_format == "odt":
                # Already ODT, just copy
                shutil.copy2(input_path, output_path)
                return True, ""

            if self.dependencies["pandoc"]:
                return self._convert_to_odt_pandoc(
                    input_path, output_path, input_format, options
                )
            else:
                return False, "ODT conversion requires Pandoc"

        except Exception as e:
            return False, f"ODT conversion error: {str(e)}"

    def _convert_to_odt_pandoc(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert to ODT using Pandoc"""
        try:
            format_map = {
                "docx": "docx",
                # Note: DOC format not supported by Pandoc
                "html": "html",
                "markdown": "markdown",
                "txt": "plain",
                "rtf": "rtf",
            }

            pandoc_input_format = format_map.get(input_format, "plain")

            pypandoc.convert_file(
                input_path, "odt", format=pandoc_input_format, outputfile=output_path
            )

            return True, ""

        except Exception as e:
            logger.error(f"Pandoc ODT conversion error: {e}")
            return False, str(e)

    def _convert_to_txt(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert document to plain text"""
        try:
            if input_format == "txt":
                # Already TXT, just copy
                shutil.copy2(input_path, output_path)
                return True, ""

            # Route to appropriate text extraction method
            if input_format == "pdf":
                return self._extract_text_from_pdf(input_path, output_path, options)
            elif input_format == "docx" and self.dependencies["python_docx"]:
                return self._extract_text_from_docx(input_path, output_path, options)
            elif input_format == "doc":
                return self._extract_text_from_doc(input_path, output_path, options)
            elif input_format == "html":
                return self._extract_text_from_html(input_path, output_path, options)
            elif input_format == "markdown":
                return self._extract_text_from_markdown(
                    input_path, output_path, options
                )
            elif self.dependencies["pandoc"]:
                return self._convert_to_txt_pandoc(
                    input_path, output_path, input_format, options
                )
            else:
                return False, f"No text extraction method available for {input_format}"

        except Exception as e:
            return False, f"Text conversion error: {str(e)}"

    def _extract_text_from_pdf(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Extract text from PDF"""
        try:
            extracted_text = ""

            if self.dependencies["pdf_processing"]:
                # Use pdfplumber for text extraction
                import pdfplumber

                with pdfplumber.open(input_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            extracted_text += text + "\n"

            # OCR fallback if enabled and no text found
            if (
                not extracted_text.strip()
                and options.get("ocr_scan")
                and self.dependencies["ocr"]
            ):
                try:
                    import fitz  # PyMuPDF

                    pdf_document = fitz.open(input_path)
                    for page_num in range(pdf_document.page_count):
                        page = pdf_document.load_page(page_num)
                        pix = page.get_pixmap()
                        img = PILImage.frombytes(
                            "RGB", [pix.width, pix.height], pix.samples
                        )
                        text = pytesseract.image_to_string(img)
                        extracted_text += text + "\n"
                    pdf_document.close()
                except Exception as ocr_error:
                    logger.warning(f"OCR extraction failed: {ocr_error}")

            if not extracted_text.strip():
                return False, "No text could be extracted from PDF"

            # Write extracted text
            with open(
                output_path, "w", encoding=options.get("text_encoding", "utf-8")
            ) as f:
                f.write(extracted_text)

            return True, ""

        except Exception as e:
            logger.error(f"PDF text extraction error: {e}")
            return False, str(e)

    def _extract_text_from_docx(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Extract text from DOCX"""
        try:
            doc = DocxDocument(input_path)

            extracted_text = []
            for paragraph in doc.paragraphs:
                extracted_text.append(paragraph.text)

            # Write extracted text
            with open(
                output_path, "w", encoding=options.get("text_encoding", "utf-8")
            ) as f:
                f.write("\n".join(extracted_text))

            return True, ""

        except Exception as e:
            logger.error(f"DOCX text extraction error: {e}")
            return False, str(e)

    def _extract_text_from_doc(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Extract text from DOC (old Word format)"""
        try:
            # DOC files are proprietary binary format - try multiple approaches

            # Method 1: Windows COM automation (best for DOC files on Windows)
            if self.dependencies["win32com"]:
                try:
                    extracted_text = self._extract_doc_with_com(input_path)
                    if extracted_text:
                        with open(
                            output_path,
                            "w",
                            encoding=options.get("text_encoding", "utf-8"),
                        ) as f:
                            f.write(extracted_text)
                        return True, ""
                    else:
                        logger.warning("COM automation extracted no text from DOC file")
                except Exception as com_error:
                    logger.warning(f"COM automation failed: {com_error}")

            # Method 2: Try docx2txt library if available
            if self.dependencies["docx2txt"]:
                try:
                    import docx2txt

                    extracted_text = docx2txt.process(input_path)

                    if extracted_text:
                        with open(
                            output_path,
                            "w",
                            encoding=options.get("text_encoding", "utf-8"),
                        ) as f:
                            f.write(extracted_text)
                        return True, ""
                    else:
                        logger.warning("docx2txt extracted no text from DOC file")

                except Exception as docx2txt_error:
                    logger.warning(f"docx2txt failed: {docx2txt_error}")

            # Method 3: Fallback - try python-docx (rarely works for true DOC files)
            if self.dependencies["python_docx"]:
                try:
                    doc = DocxDocument(input_path)
                    extracted_text = []
                    for paragraph in doc.paragraphs:
                        extracted_text.append(paragraph.text)

                    if extracted_text:
                        with open(
                            output_path,
                            "w",
                            encoding=options.get("text_encoding", "utf-8"),
                        ) as f:
                            f.write("\n".join(extracted_text))
                        return True, ""

                except Exception as docx_error:
                    logger.warning(f"python-docx failed for DOC file: {docx_error}")

            # If all methods fail, provide helpful error message
            return False, (
                "DOC format (Word 97-2003) conversion failed. "
                "This appears to be a genuine legacy DOC file that requires Microsoft Word to read properly. "
                "Please convert your DOC file to DOCX format using Microsoft Word, LibreOffice, or Google Docs, "
                "then try the conversion again. Alternatively, save as RTF or TXT format for easier conversion."
            )

        except Exception as e:
            logger.error(f"DOC text extraction error: {e}")
            return False, str(e)

    def _extract_doc_with_com(self, input_path: str) -> str:
        """Extract text from DOC file using Windows COM automation"""
        try:
            import os

            import pythoncom
            import win32com.client

            # Initialize COM
            pythoncom.CoInitialize()

            # Create Word application
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False

            try:
                # Open the document
                abs_path = os.path.abspath(input_path)
                doc = word_app.Documents.Open(abs_path, ReadOnly=True)

                # Extract text content
                extracted_text = doc.Content.Text

                # Close the document
                doc.Close(SaveChanges=False)

                return extracted_text

            finally:
                # Always quit Word application
                word_app.Quit()
                pythoncom.CoUninitialize()

        except Exception as e:
            logger.error(f"COM automation error: {e}")
            try:
                pythoncom.CoUninitialize()
            except:
                pass
            raise e

    def _convert_doc_via_docx(
        self,
        input_path: str,
        output_path: str,
        output_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert DOC file via DOCX intermediate format for better accuracy"""
        try:
            # Step 1: Convert DOC to DOCX using Windows COM automation
            temp_docx = tempfile.NamedTemporaryFile(
                mode="w+b", suffix=".docx", delete=False
            )
            temp_docx.close()

            success, error = self._convert_doc_to_docx_com(
                input_path, temp_docx.name, options
            )
            if not success:
                os.unlink(temp_docx.name)
                logger.warning(f"COM conversion failed: {error}")
                # Fallback to text extraction method
                return self._convert_doc_fallback(
                    input_path, output_path, output_format, options
                )

            # Step 2: Convert DOCX to target format
            try:
                if output_format == "pdf":
                    result = self._convert_to_pdf(
                        temp_docx.name, output_path, "docx", options
                    )
                elif output_format == "docx":
                    # Already DOCX, just copy
                    shutil.copy2(temp_docx.name, output_path)
                    result = (True, "")
                elif output_format == "html":
                    result = self._convert_to_html(
                        temp_docx.name, output_path, "docx", options
                    )
                elif output_format == "markdown":
                    result = self._convert_to_markdown(
                        temp_docx.name, output_path, "docx", options
                    )
                elif output_format == "rtf":
                    result = self._convert_to_rtf(
                        temp_docx.name, output_path, "docx", options
                    )
                elif output_format == "odt":
                    result = self._convert_to_odt(
                        temp_docx.name, output_path, "docx", options
                    )
                else:
                    result = (False, f"Unsupported output format: {output_format}")

                # Clean up temporary DOCX
                os.unlink(temp_docx.name)

                return result

            except Exception as conversion_error:
                os.unlink(temp_docx.name)
                logger.error(
                    f"DOCX to {output_format} conversion failed: {conversion_error}"
                )
                return (
                    False,
                    f"Conversion from DOCX to {output_format} failed: {str(conversion_error)}",
                )

        except Exception as e:
            logger.error(f"DOC via DOCX conversion error: {e}")
            return False, f"DOC conversion error: {str(e)}"

    def _convert_doc_to_docx_com(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Convert DOC to DOCX using Windows COM automation"""
        if not self.dependencies["win32com"]:
            return False, "Windows COM not available"

        try:
            import os

            import pythoncom
            import win32com.client

            # Initialize COM
            pythoncom.CoInitialize()

            # Create Word application
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False

            try:
                # Open the DOC document
                abs_input_path = os.path.abspath(input_path)
                abs_output_path = os.path.abspath(output_path)

                doc = word_app.Documents.Open(abs_input_path, ReadOnly=True)

                # Save as DOCX format (format code 16 = docx)
                doc.SaveAs2(abs_output_path, FileFormat=16)

                # Close the document
                doc.Close(SaveChanges=False)

                return True, ""

            finally:
                # Always quit Word application
                word_app.Quit()
                pythoncom.CoUninitialize()

        except Exception as e:
            logger.error(f"COM DOC to DOCX conversion error: {e}")
            try:
                pythoncom.CoUninitialize()
            except:
                pass
            return False, str(e)

    def _convert_doc_fallback(
        self,
        input_path: str,
        output_path: str,
        output_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Fallback DOC conversion using text extraction"""
        try:
            # Extract text first
            temp_txt = tempfile.NamedTemporaryFile(
                mode="w+", suffix=".txt", delete=False
            )
            temp_txt.close()

            success, error = self._extract_text_from_doc(
                input_path, temp_txt.name, options
            )
            if not success:
                os.unlink(temp_txt.name)
                return False, f"Text extraction failed: {error}"

            # Convert text to target format
            try:
                if output_format == "pdf":
                    result = self._convert_to_pdf(
                        temp_txt.name, output_path, "txt", options
                    )
                elif output_format == "html":
                    result = self._convert_to_html(
                        temp_txt.name, output_path, "txt", options
                    )
                elif output_format == "markdown":
                    result = self._convert_to_markdown(
                        temp_txt.name, output_path, "txt", options
                    )
                elif output_format == "rtf":
                    result = self._convert_to_rtf(
                        temp_txt.name, output_path, "txt", options
                    )
                elif output_format == "docx":
                    result = self._convert_to_docx(
                        temp_txt.name, output_path, "txt", options
                    )
                elif output_format == "odt":
                    result = self._convert_to_odt(
                        temp_txt.name, output_path, "txt", options
                    )
                else:
                    result = (False, f"Unsupported fallback format: {output_format}")

                os.unlink(temp_txt.name)
                return result

            except Exception as conversion_error:
                os.unlink(temp_txt.name)
                return False, f"Fallback conversion failed: {str(conversion_error)}"

        except Exception as e:
            return False, f"DOC fallback conversion error: {str(e)}"

    def _convert_docx_to_pdf_com(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Convert DOCX to PDF using Windows COM automation (best formatting preservation)"""
        if not self.dependencies["win32com"]:
            return False, "Windows COM not available"

        try:
            import os

            import pythoncom
            import win32com.client

            # Initialize COM
            pythoncom.CoInitialize()

            # Create Word application
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False

            try:
                # Open the DOCX document
                abs_input_path = os.path.abspath(input_path)
                abs_output_path = os.path.abspath(output_path)

                doc = word_app.Documents.Open(abs_input_path, ReadOnly=True)

                # Configure PDF export options
                # PDF format constant
                wdExportFormatPDF = 17
                wdExportOptimizeForPrint = 0

                # Export as PDF with error handling for different Word versions
                try:
                    # Try full parameter set first
                    doc.ExportAsFixedFormat(
                        OutputFileName=abs_output_path,
                        ExportFormat=wdExportFormatPDF,
                        OpenAfterExport=False,
                        OptimizeFor=wdExportOptimizeForPrint,
                        BitmapMissingFonts=True,
                        DocStructureTags=True,
                        CreateBookmarks=True,
                    )
                except Exception as export_error:
                    logger.warning(f"Full parameter export failed: {export_error}")
                    # Try minimal parameter set for older Word versions
                    doc.ExportAsFixedFormat(
                        OutputFileName=abs_output_path, ExportFormat=wdExportFormatPDF
                    )

                # Close the document
                doc.Close(SaveChanges=False)

                return True, ""

            finally:
                # Always quit Word application
                word_app.Quit()
                pythoncom.CoUninitialize()

        except Exception as e:
            logger.error(f"COM DOCX to PDF conversion error: {e}")
            try:
                pythoncom.CoUninitialize()
            except:
                pass
            return False, str(e)

    def _convert_docx_to_html_com(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Convert DOCX to HTML using Windows COM automation (best quality)"""
        if not self.dependencies["win32com"]:
            return False, "Windows COM not available"

        try:
            import os

            import pythoncom
            import win32com.client

            # Initialize COM
            pythoncom.CoInitialize()

            # Create Word application
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False

            try:
                # Open the DOCX document
                abs_input_path = os.path.abspath(input_path)
                abs_output_path = os.path.abspath(output_path)

                doc = word_app.Documents.Open(abs_input_path, ReadOnly=True)

                # Save as filtered HTML (format code 10)
                # This preserves more formatting than basic HTML
                wdFormatFilteredHTML = 10

                doc.SaveAs2(
                    FileName=abs_output_path,
                    FileFormat=wdFormatFilteredHTML,
                    EmbedTrueTypeFonts=False,
                    SaveNativePictureFormat=options.get("include_images", True),
                    SaveFormsData=options.get("include_tables", True),
                )

                # Close the document
                doc.Close(SaveChanges=False)

                # Enhance the generated HTML
                self._enhance_com_generated_html(abs_output_path, options)

                return True, ""

            finally:
                # Always quit Word application
                word_app.Quit()
                pythoncom.CoUninitialize()

        except Exception as e:
            logger.error(f"COM DOCX to HTML conversion error: {e}")
            try:
                pythoncom.CoUninitialize()
            except:
                pass
            return False, str(e)

    def _enhance_com_generated_html(self, html_path: str, options: Dict[str, Any]):
        """Enhance COM-generated HTML with better styling and structure"""
        try:
            with open(html_path, "r", encoding="utf-8", errors="ignore") as f:
                html_content = f.read()

            if self.dependencies["bs4"]:
                soup = BeautifulSoup(html_content, "html.parser")

                # Add modern HTML5 doctype if missing
                if not html_content.strip().startswith("<!DOCTYPE"):
                    soup.insert(0, BeautifulSoup("<!DOCTYPE html>", "html.parser"))

                # Enhance head section
                head = soup.find("head")
                if not head:
                    head = soup.new_tag("head")
                    soup.html.insert(0, head)

                # Add viewport meta tag
                if not soup.find("meta", {"name": "viewport"}):
                    viewport_meta = soup.new_tag(
                        "meta",
                        name="viewport",
                        content="width=device-width, initial-scale=1.0",
                    )
                    head.insert(0, viewport_meta)

                # Add enhanced styling
                enhanced_css = self._get_enhanced_html_css(options)
                style_tag = soup.new_tag("style")
                style_tag.string = enhanced_css
                head.append(style_tag)

                # Clean up Word-specific elements
                for element in soup.find_all(
                    attrs={
                        "class": lambda x: x and ("Mso" in str(x) or "mso" in str(x))
                    }
                ):
                    element.extract()

                # Enhance table formatting
                for table in soup.find_all("table"):
                    table["class"] = "enhanced-table"

                # Enhance headings
                for i in range(1, 7):
                    for heading in soup.find_all(f"h{i}"):
                        heading["class"] = f"enhanced-heading level-{i}"

                # Write enhanced HTML
                with open(html_path, "w", encoding="utf-8") as f:
                    f.write(str(soup))

        except Exception as e:
            logger.warning(f"HTML enhancement failed: {e}")

    def _get_enhanced_html_css(self, options: Dict[str, Any]) -> str:
        """Get enhanced CSS for better HTML presentation"""
        return """
/* Enhanced Document Styles */
body {
    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    line-height: 1.6;
    max-width: 1200px;
    margin: 0 auto;
    padding: 20px;
    background-color: #ffffff;
    color: #333333;
}

/* Enhanced Headings */
.enhanced-heading {
    margin-top: 2em;
    margin-bottom: 1em;
    border-bottom: 2px solid #e9ecef;
    padding-bottom: 0.5em;
}

.enhanced-heading.level-1 { font-size: 2.5em; color: #2c3e50; }
.enhanced-heading.level-2 { font-size: 2em; color: #34495e; }
.enhanced-heading.level-3 { font-size: 1.75em; color: #7f8c8d; }
.enhanced-heading.level-4 { font-size: 1.5em; color: #95a5a6; }
.enhanced-heading.level-5 { font-size: 1.25em; color: #bdc3c7; }
.enhanced-heading.level-6 { font-size: 1.1em; color: #ecf0f1; }

/* Enhanced Tables */
.enhanced-table {
    width: 100%;
    border-collapse: collapse;
    margin: 20px 0;
    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    border-radius: 8px;
    overflow: hidden;
}

.enhanced-table th,
.enhanced-table td {
    padding: 12px 15px;
    text-align: left;
    border-bottom: 1px solid #ddd;
}

.enhanced-table th {
    background-color: #3498db;
    color: white;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.5px;
}

.enhanced-table tr:nth-child(even) {
    background-color: #f8f9fa;
}

.enhanced-table tr:hover {
    background-color: #e3f2fd;
    transform: scale(1.01);
    transition: all 0.3s ease;
}

/* Enhanced Paragraphs */
p {
    margin-bottom: 1em;
    text-align: justify;
    hyphens: auto;
}

/* Enhanced Lists */
ul, ol {
    padding-left: 30px;
    margin-bottom: 1em;
}

li {
    margin-bottom: 0.5em;
}

/* Enhanced Links */
a {
    color: #3498db;
    text-decoration: none;
    border-bottom: 1px dotted #3498db;
    transition: all 0.3s ease;
}

a:hover {
    color: #2980b9;
    border-bottom: 1px solid #2980b9;
}

/* Enhanced Images */
img {
    max-width: 100%;
    height: auto;
    border-radius: 8px;
    box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    margin: 10px 0;
}

/* Enhanced Blockquotes */
blockquote {
    border-left: 4px solid #3498db;
    margin: 20px 0;
    padding: 15px 20px;
    background-color: #f8f9fa;
    font-style: italic;
    border-radius: 0 8px 8px 0;
}

/* Enhanced Code */
code {
    background-color: #f1f2f6;
    padding: 2px 6px;
    border-radius: 4px;
    font-family: 'Monaco', 'Menlo', 'Ubuntu Mono', monospace;
    font-size: 0.9em;
}

pre {
    background-color: #2f3640;
    color: #f5f6fa;
    padding: 20px;
    border-radius: 8px;
    overflow-x: auto;
    margin: 20px 0;
}

/* Responsive Design */
@media (max-width: 768px) {
    body {
        padding: 10px;
    }

    .enhanced-heading.level-1 { font-size: 2em; }
    .enhanced-heading.level-2 { font-size: 1.75em; }

    .enhanced-table {
        font-size: 0.9em;
    }

    .enhanced-table th,
    .enhanced-table td {
        padding: 8px 10px;
    }
}

/* Print Styles */
@media print {
    body {
        color: black;
        background: white;
    }

    .enhanced-table {
        box-shadow: none;
    }

    .enhanced-table tr:hover {
        background-color: transparent;
        transform: none;
    }
}
"""

    def _extract_text_from_html(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Extract text from HTML"""
        try:
            # Use robust encoding detection
            html_content, encoding_used = self._read_file_with_encoding_detection(
                input_path
            )
            logger.debug(
                f"Read HTML file for text extraction with encoding: {encoding_used}"
            )

            if self.dependencies["bs4"]:
                soup = BeautifulSoup(html_content, "html.parser")
                extracted_text = soup.get_text(separator="\n")
            else:
                # Fallback: simple HTML tag removal
                import re

                extracted_text = re.sub(r"<[^>]+>", "", html_content)

            # Write extracted text
            with open(
                output_path, "w", encoding=options.get("text_encoding", "utf-8")
            ) as f:
                f.write(extracted_text)

            return True, ""

        except Exception as e:
            logger.error(f"HTML text extraction error: {e}")
            return False, str(e)

    def _extract_text_from_markdown(
        self, input_path: str, output_path: str, options: Dict[str, Any]
    ) -> Tuple[bool, str]:
        """Extract text from Markdown"""
        try:
            # Use robust encoding detection
            md_content, encoding_used = self._read_file_with_encoding_detection(
                input_path
            )
            logger.debug(
                f"Read Markdown file for text extraction with encoding: {encoding_used}"
            )

            if self.dependencies["markdown"]:
                # Convert to HTML first, then extract text
                md = markdown.Markdown()
                html = md.convert(md_content)

                if self.dependencies["bs4"]:
                    soup = BeautifulSoup(html, "html.parser")
                    extracted_text = soup.get_text(separator="\n")
                else:
                    # Simple markdown stripping
                    import re

                    # Remove markdown syntax
                    text = re.sub(r"#+\s*", "", md_content)  # Headers
                    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # Bold
                    text = re.sub(r"\*(.*?)\*", r"\1", text)  # Italic
                    text = re.sub(r"`(.*?)`", r"\1", text)  # Inline code
                    extracted_text = text
            else:
                # Fallback: treat as plain text
                extracted_text = md_content

            # Write extracted text
            with open(
                output_path, "w", encoding=options.get("text_encoding", "utf-8")
            ) as f:
                f.write(extracted_text)

            return True, ""

        except Exception as e:
            logger.error(f"Markdown text extraction error: {e}")
            return False, str(e)

    def _convert_to_txt_pandoc(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        options: Dict[str, Any],
    ) -> Tuple[bool, str]:
        """Convert to TXT using Pandoc"""
        try:
            format_map = {
                "docx": "docx",
                # Note: DOC format not supported by Pandoc
                "html": "html",
                "markdown": "markdown",
                "rtf": "rtf",
                "odt": "odt",
                "epub": "epub",
            }

            pandoc_input_format = format_map.get(input_format, "plain")

            pypandoc.convert_file(
                input_path, "plain", format=pandoc_input_format, outputfile=output_path
            )

            return True, ""

        except Exception as e:
            logger.error(f"Pandoc TXT conversion error: {e}")
            return False, str(e)
