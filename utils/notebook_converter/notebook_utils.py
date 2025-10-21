# utils/notebook_converter/notebook_utils.py - NOTEBOOK CONVERTER UTILITIES
# Dedicated utilities for Jupyter notebook conversion
import os
import tempfile
import shutil
import json
from pathlib import Path
import logging
from typing import List, Optional, Tuple, Dict, Any

# Universal Security Framework Integration
try:
    from security.core.validators import validate_content, validate_filename
    from security.core.sanitizers import sanitize_filename
    SECURITY_FRAMEWORK_AVAILABLE = True
except ImportError:
    SECURITY_FRAMEWORK_AVAILABLE = False

# Jupyter notebook conversion dependencies
try:
    import nbformat
    from nbconvert import HTMLExporter, PDFExporter, MarkdownExporter
    from nbconvert.writers import FilesWriter
    from nbconvert.preprocessors import ExecutePreprocessor, ClearOutputPreprocessor
    NBCONVERT_AVAILABLE = True
except ImportError:
    NBCONVERT_AVAILABLE = False

# Additional dependencies for advanced conversions
try:
    import pandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import docx.shared
    import docx.enum.style
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError) as e:
    WEASYPRINT_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

# Initialize logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class NotebookConverter:
    """
    Dedicated Jupyter notebook converter supporting:
    - Notebook to PDF conversion
    - Notebook to HTML conversion
    - Notebook to DOCX conversion  
    - Notebook to Markdown conversion
    """
    
    def __init__(self):
        self.temp_dirs = []  # Track temporary directories for cleanup
        
        # Check for required dependencies
        self.dependencies = self._check_dependencies()
        
    def __del__(self):
        """Cleanup temporary directories on object destruction"""
        self.cleanup_temp_dirs()
        
    def cleanup_temp_dirs(self):
        """Clean up all temporary directories created by this instance"""
        for temp_dir in self.temp_dirs:
            if os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                except Exception as e:
                    logger.warning(f"Failed to cleanup temp dir {temp_dir}: {e}")
        self.temp_dirs.clear()
    
    def _check_dependencies(self) -> Dict[str, bool]:
        """Check availability of all dependencies"""
        return {
            'nbconvert': NBCONVERT_AVAILABLE,
            'pandoc': PANDOC_AVAILABLE,
            'python_docx': PYTHON_DOCX_AVAILABLE,
            'weasyprint': WEASYPRINT_AVAILABLE,
            'bs4': BS4_AVAILABLE
        }
    
    def is_pdf_conversion_available(self) -> bool:
        """Check if PDF conversion is available"""
        return self.dependencies['nbconvert']
    
    def is_html_conversion_available(self) -> bool:
        """Check if HTML conversion is available"""
        return self.dependencies['nbconvert']
    
    def is_docx_conversion_available(self) -> bool:
        """Check if DOCX conversion is available"""
        return (self.dependencies['nbconvert'] and 
                self.dependencies['python_docx'])
    
    def is_markdown_conversion_available(self) -> bool:
        """Check if Markdown conversion is available"""
        return self.dependencies['nbconvert']
    
    def validate_file_security(self, file_path: str) -> Tuple[bool, List[str]]:
        """Validate file using universal security framework"""
        if not SECURITY_FRAMEWORK_AVAILABLE:
            logger.warning("Security framework not available - performing basic validation")
            return self._basic_file_validation(file_path)
        
        try:
            # Read file content for security validation
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            # Determine file type
            file_ext = Path(file_path).suffix.lower().lstrip('.')
            
            # Perform security validation
            is_safe, issues = validate_content(file_content, file_ext)
            
            if not is_safe:
                logger.error(f"Security validation failed for {file_path}: {issues}")
                return False, issues
            
            logger.info(f"Security validation passed for {file_path}")
            return True, []
            
        except Exception as e:
            logger.error(f"Security validation error for {file_path}: {e}")
            return False, [f"Security validation error: {str(e)}"]
    
    def _basic_file_validation(self, file_path: str) -> Tuple[bool, List[str]]:
        """Basic file validation when security framework is not available"""
        try:
            if not os.path.exists(file_path):
                return False, ["File does not exist"]
            
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                return False, ["File is empty"]
            
            if file_size > 100 * 1024 * 1024:  # 100MB limit
                return False, ["File too large"]
            
            # Validate JSON structure for notebook
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    notebook_data = json.load(f)
                
                # Basic notebook structure validation
                if 'cells' not in notebook_data:
                    return False, ["Invalid notebook format: missing 'cells'"]
                
                if 'metadata' not in notebook_data:
                    return False, ["Invalid notebook format: missing 'metadata'"]
                
                if 'nbformat' not in notebook_data:
                    return False, ["Invalid notebook format: missing 'nbformat'"]
                    
            except json.JSONDecodeError as e:
                return False, [f"Invalid JSON format: {str(e)}"]
            
            return True, []
            
        except Exception as e:
            return False, [f"Basic validation error: {str(e)}"]
    
    def _load_notebook(self, input_path: str) -> Optional[nbformat.NotebookNode]:
        """Load and validate notebook file"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                notebook = nbformat.read(f, as_version=4)
            
            # Validate notebook
            nbformat.validate(notebook)
            return notebook
            
        except Exception as e:
            logger.error(f"Failed to load notebook {input_path}: {e}")
            return None
    
    def _preprocess_notebook(self, notebook: nbformat.NotebookNode, 
                           include_code_cells: bool = True,
                           include_outputs: bool = True,
                           include_metadata: bool = True) -> nbformat.NotebookNode:
        """Preprocess notebook based on options"""
        try:
            # Create a copy to avoid modifying original
            processed_notebook = nbformat.v4.new_notebook()
            processed_notebook.metadata = notebook.metadata.copy() if include_metadata else {}
            
            for cell in notebook.cells:
                # Handle code cells
                if cell.cell_type == 'code':
                    if include_code_cells:
                        new_cell = nbformat.v4.new_code_cell(
                            source=cell.source,
                            metadata=cell.metadata if include_metadata else {}
                        )
                        if include_outputs and hasattr(cell, 'outputs'):
                            new_cell.outputs = cell.outputs
                        else:
                            new_cell.outputs = []
                        processed_notebook.cells.append(new_cell)
                
                # Handle markdown cells
                elif cell.cell_type == 'markdown':
                    new_cell = nbformat.v4.new_markdown_cell(
                        source=cell.source,
                        metadata=cell.metadata if include_metadata else {}
                    )
                    processed_notebook.cells.append(new_cell)
                
                # Handle raw cells
                elif cell.cell_type == 'raw':
                    new_cell = nbformat.v4.new_raw_cell(
                        source=cell.source,
                        metadata=cell.metadata if include_metadata else {}
                    )
                    processed_notebook.cells.append(new_cell)
            
            return processed_notebook
            
        except Exception as e:
            logger.error(f"Failed to preprocess notebook: {e}")
            return notebook
    
    def notebook_to_pdf(self, input_path: str, output_path: str,
                       include_code_cells: bool = True,
                       include_outputs: bool = True,
                       include_metadata: bool = True,
                       include_images: bool = True) -> bool:
        """Convert notebook to PDF"""
        try:
            if not self.is_pdf_conversion_available():
                logger.error("Required dependencies not available for PDF conversion")
                return False
            
            if not os.path.exists(input_path):
                logger.error(f"Input file not found: {input_path}")
                return False
            
            # Security validation
            is_safe, security_issues = self.validate_file_security(input_path)
            if not is_safe:
                logger.error(f"Security validation failed for {input_path}: {security_issues}")
                return False
            
            logger.info(f"Converting notebook to PDF: {input_path} -> {output_path}")
            
            # Load notebook
            notebook = self._load_notebook(input_path)
            if not notebook:
                return False
            
            # Preprocess notebook
            processed_notebook = self._preprocess_notebook(
                notebook, include_code_cells, include_outputs, include_metadata
            )
            
            # Create PDF exporter
            try:
                pdf_exporter = PDFExporter()
                
                # Configure exporter
                pdf_exporter.template_name = 'classic'
                
                # Convert to PDF
                (body, resources) = pdf_exporter.from_notebook_node(processed_notebook)
                
                # Write PDF file
                with open(output_path, 'wb') as f:
                    f.write(body)
                    
            except Exception as pdf_error:
                logger.error(f"PDF exporter error: {pdf_error}")
                # Try alternative PDF conversion approach
                try:
                    # Fallback: convert to HTML first, then to PDF if possible
                    html_exporter = HTMLExporter()
                    (html_body, html_resources) = html_exporter.from_notebook_node(processed_notebook)
                    
                    # Save HTML temporarily
                    html_temp = output_path.replace('.pdf', '_temp.html')
                    with open(html_temp, 'w', encoding='utf-8') as f:
                        f.write(html_body)
                    
                    logger.info(f"PDF conversion failed, created HTML instead: {html_temp}")
                    # For now, return False to indicate PDF conversion failed
                    return False
                    
                except Exception as fallback_error:
                    logger.error(f"PDF fallback conversion also failed: {fallback_error}")
                    return False
                
            success = os.path.exists(output_path)
            if success:
                logger.info(f"PDF conversion successful: {output_path}")
            else:
                logger.error("PDF conversion failed - output file not created")
                
            return success
            
        except Exception as e:
            logger.error(f"PDF conversion error: {e}")
            return False
    
    def notebook_to_html(self, input_path: str, output_path: str,
                        include_code_cells: bool = True,
                        include_outputs: bool = True,
                        include_metadata: bool = True,
                        include_images: bool = True) -> bool:
        """Convert notebook to HTML"""
        try:
            if not self.is_html_conversion_available():
                logger.error("Required dependencies not available for HTML conversion")
                return False
            
            if not os.path.exists(input_path):
                logger.error(f"Input file not found: {input_path}")
                return False
            
            # Security validation
            is_safe, security_issues = self.validate_file_security(input_path)
            if not is_safe:
                logger.error(f"Security validation failed for {input_path}: {security_issues}")
                return False
            
            logger.info(f"Converting notebook to HTML: {input_path} -> {output_path}")
            
            # Load notebook
            notebook = self._load_notebook(input_path)
            if not notebook:
                logger.error("Failed to load notebook file")
                return False
            
            # Preprocess notebook
            processed_notebook = self._preprocess_notebook(
                notebook, include_code_cells, include_outputs, include_metadata
            )
            
            # Create HTML exporter
            try:
                html_exporter = HTMLExporter()
                # Configure exporter
                html_exporter.template_name = 'classic'
            except Exception as exporter_error:
                logger.error(f"Failed to create HTML exporter: {exporter_error}")
                return False
            
            try:
                # Convert to HTML
                (body, resources) = html_exporter.from_notebook_node(processed_notebook)
            except Exception as conversion_error:
                logger.error(f"Failed to convert notebook to HTML: {conversion_error}")
                return False
            
            try:
                # Write HTML file
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(body)
            except Exception as file_error:
                logger.error(f"Failed to write HTML file: {file_error}")
                return False
            
            success = os.path.exists(output_path)
            if success:
                logger.info(f"HTML conversion successful: {output_path}")
            else:
                logger.error("HTML conversion failed - output file not created")
                
            return success
            
        except Exception as e:
            logger.error(f"HTML conversion error: {e}")
            return False
    
    def notebook_to_markdown(self, input_path: str, output_path: str,
                           include_code_cells: bool = True,
                           include_outputs: bool = True,
                           include_metadata: bool = True,
                           include_images: bool = True) -> bool:
        """Convert notebook to Markdown"""
        try:
            if not self.is_markdown_conversion_available():
                logger.error("Required dependencies not available for Markdown conversion")
                return False
            
            if not os.path.exists(input_path):
                logger.error(f"Input file not found: {input_path}")
                return False
            
            # Security validation
            is_safe, security_issues = self.validate_file_security(input_path)
            if not is_safe:
                logger.error(f"Security validation failed for {input_path}: {security_issues}")
                return False
            
            logger.info(f"Converting notebook to Markdown: {input_path} -> {output_path}")
            
            # Load notebook
            notebook = self._load_notebook(input_path)
            if not notebook:
                return False
            
            # Preprocess notebook
            processed_notebook = self._preprocess_notebook(
                notebook, include_code_cells, include_outputs, include_metadata
            )
            
            # Create Markdown exporter
            md_exporter = MarkdownExporter()
            
            # Convert to Markdown
            (body, resources) = md_exporter.from_notebook_node(processed_notebook)
            
            # Write Markdown file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(body)
            
            success = os.path.exists(output_path)
            if success:
                logger.info(f"Markdown conversion successful: {output_path}")
            else:
                logger.error("Markdown conversion failed - output file not created")
                
            return success
            
        except Exception as e:
            logger.error(f"Markdown conversion error: {e}")
            return False
    
    def notebook_to_docx(self, input_path: str, output_path: str,
                        include_code_cells: bool = True,
                        include_outputs: bool = True,
                        include_metadata: bool = True,
                        include_images: bool = True) -> bool:
        """Convert notebook to DOCX via HTML intermediate"""
        try:
            if not self.is_docx_conversion_available():
                logger.error("Required dependencies not available for DOCX conversion")
                return False
            
            if not os.path.exists(input_path):
                logger.error(f"Input file not found: {input_path}")
                return False
            
            # Security validation
            is_safe, security_issues = self.validate_file_security(input_path)
            if not is_safe:
                logger.error(f"Security validation failed for {input_path}: {security_issues}")
                return False
            
            logger.info(f"Converting notebook to DOCX: {input_path} -> {output_path}")
            
            # Load notebook
            notebook = self._load_notebook(input_path)
            if not notebook:
                return False
            
            # Preprocess notebook
            processed_notebook = self._preprocess_notebook(
                notebook, include_code_cells, include_outputs, include_metadata
            )
            
            # Create HTML exporter first
            html_exporter = HTMLExporter()
            html_exporter.template_name = 'basic'
            
            # Convert to HTML
            (html_body, resources) = html_exporter.from_notebook_node(processed_notebook)
            
            # Create DOCX document
            doc = Document()
            
            # Parse HTML and convert to DOCX
            if BS4_AVAILABLE:
                soup = BeautifulSoup(html_body, 'html.parser')
                self._html_to_docx(soup, doc, include_code_cells)
            else:
                # Fallback: basic text extraction
                self._notebook_to_docx_basic(processed_notebook, doc, include_code_cells)
            
            # Save DOCX file
            doc.save(output_path)
            
            success = os.path.exists(output_path)
            if success:
                logger.info(f"DOCX conversion successful: {output_path}")
            else:
                logger.error("DOCX conversion failed - output file not created")
                
            return success
            
        except Exception as e:
            logger.error(f"DOCX conversion error: {e}")
            return False
    
    def _html_to_docx(self, soup: BeautifulSoup, doc: Document, include_code_cells: bool):
        """Convert HTML BeautifulSoup to DOCX document"""
        try:
            # Check if the document has 'Code' and 'Quote' styles, if not, add them
            self._ensure_docx_styles(doc)
            
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'div']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # Handle headers
                    heading = doc.add_heading(element.get_text().strip(), level=int(element.name[1]))
                elif element.name == 'pre' or 'highlight' in element.get('class', []):
                    # Handle code blocks
                    if include_code_cells:
                        code_para = doc.add_paragraph(element.get_text())
                        # Apply code style safely
                        try:
                            if 'Code' in doc.styles:
                                code_para.style = 'Code'
                        except Exception:
                            # If style application fails, continue with default style
                            pass
                elif element.name == 'p':
                    # Handle paragraphs
                    text = element.get_text().strip()
                    if text:
                        doc.add_paragraph(text)
                elif element.name == 'div' and 'output' in element.get('class', []):
                    # Handle outputs
                    output_text = element.get_text().strip()
                    if output_text:
                        output_para = doc.add_paragraph(output_text)
                        # Apply quote style safely
                        try:
                            if 'Quote' in doc.styles:
                                output_para.style = 'Quote'
                        except Exception:
                            # If style application fails, continue with default style
                            pass
        except Exception as e:
            logger.warning(f"Error converting HTML to DOCX: {e}")
            
    def _ensure_docx_styles(self, doc: Document):
        """Ensure required styles exist in the document"""
        try:
            # Add Code style if not present
            if 'Code' not in doc.styles:
                code_style = doc.styles.add_style('Code', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
                code_style.font.name = 'Courier New'
                code_style.font.size = docx.shared.Pt(9)
                code_style.paragraph_format.space_before = docx.shared.Pt(6)
                code_style.paragraph_format.space_after = docx.shared.Pt(6)
                logger.info("Added 'Code' style to DOCX document")
                
            # Add Quote style if not present
            if 'Quote' not in doc.styles:
                quote_style = doc.styles.add_style('Quote', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
                quote_style.font.italic = True
                quote_style.paragraph_format.left_indent = docx.shared.Inches(0.5)
                quote_style.paragraph_format.space_before = docx.shared.Pt(6)
                quote_style.paragraph_format.space_after = docx.shared.Pt(6)
                logger.info("Added 'Quote' style to DOCX document")
                
        except Exception as e:
            logger.warning(f"Error adding styles to DOCX document: {e}")
    
    def _notebook_to_docx_basic(self, notebook: nbformat.NotebookNode, doc: Document, include_code_cells: bool):
        """Basic notebook to DOCX conversion without HTML parsing"""
        try:
            # Ensure document has required styles
            self._ensure_docx_styles(doc)
            
            for cell in notebook.cells:
                if cell.cell_type == 'markdown':
                    # Handle markdown cells
                    lines = cell.source.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            if line.startswith('#'):
                                # Handle headers
                                level = len(line) - len(line.lstrip('#'))
                                header_text = line.lstrip('# ').strip()
                                doc.add_heading(header_text, level=min(level, 6))
                            else:
                                doc.add_paragraph(line)
                
                elif cell.cell_type == 'code' and include_code_cells:
                    # Handle code cells
                    if cell.source.strip():
                        code_para = doc.add_paragraph(cell.source)
                        # Apply code style safely
                        try:
                            if 'Code' in doc.styles:
                                code_para.style = 'Code'
                        except Exception:
                            # If style application fails, continue with default style
                            pass
                    
                    # Handle outputs
                    if hasattr(cell, 'outputs'):
                        for output in cell.outputs:
                            if 'text' in output:
                                output_para = doc.add_paragraph(output.text)
                                # Apply quote style safely
                                try:
                                    if 'Quote' in doc.styles:
                                        output_para.style = 'Quote'
                                except Exception:
                                    # If style application fails, continue with default style
                                    pass
        except Exception as e:
            logger.warning(f"Error in basic notebook to DOCX conversion: {e}")

# Convenience function for quick access
def convert_notebook(input_path: str, output_path: str, output_format: str, **kwargs) -> bool:
    """
    Convert Jupyter notebook to specified format
    
    Args:
        input_path: Input notebook file path
        output_path: Output file path
        output_format: Target format ('pdf', 'html', 'docx', 'markdown')
        **kwargs: Additional conversion parameters
        
    Returns:
        bool: True if successful
    """
    converter = NotebookConverter()
    
    input_ext = Path(input_path).suffix.lower()
    
    if input_ext == '.ipynb':
        if output_format == 'pdf':
            return converter.notebook_to_pdf(input_path, output_path, **kwargs)
        elif output_format == 'html':
            return converter.notebook_to_html(input_path, output_path, **kwargs)
        elif output_format == 'docx':
            return converter.notebook_to_docx(input_path, output_path, **kwargs)
        elif output_format == 'markdown':
            return converter.notebook_to_markdown(input_path, output_path, **kwargs)
        else:
            logger.error(f"Unsupported output format: {output_format}")
            return False
    else:
        logger.error(f"Unsupported input format: {input_ext}")
        return False