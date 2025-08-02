# routes/reverse_converter_routes.py
import os
import json
from io import BytesIO
from flask import Blueprint, render_template, request, flash, redirect, send_file, current_app
from werkzeug.utils import secure_filename
import pandas as pd

# Optional imports
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_PARSER_AVAILABLE = True
except ImportError:
    HTML_PARSER_AVAILABLE = False

reverse_converter_bp = Blueprint('reverse_converter', __name__)

@reverse_converter_bp.route('/convert-to-pdf', methods=['GET', 'POST'])
def convert_to_pdf():
    """Convert various formats TO PDF"""
    
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part', 'error')
            return redirect(request.url)
        
        file = request.files['file']
        if file.filename == '':
            flash('No selected file', 'error')
            return redirect(request.url)
        
        if file:
            filename = secure_filename(file.filename)
            filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
            
            try:
                if not REPORTLAB_AVAILABLE:
                    flash('PDF conversion not supported. Please install reportlab library.', 'error')
                    return redirect(request.url)
                
                pdf_buffer = convert_file_to_pdf(filepath, filename, file_ext)
                
                if pdf_buffer:
                    flash('Successfully converted to PDF!', 'success')
                    return send_file(
                        pdf_buffer, 
                        as_attachment=True, 
                        download_name=f"{filename.rsplit('.', 1)[0]}.pdf"
                    )
                else:
                    flash(f'Unsupported file format: {file_ext}', 'error')
                    return redirect(request.url)
                    
            except Exception as e:
                current_app.logger.error(f"Error converting to PDF: {e}")
                flash('Conversion failed. Please try again.', 'error')
                return redirect(request.url)
    
    return render_template('convert_to_pdf.html')

@reverse_converter_bp.route('/convert-to-docx', methods=['GET', 'POST'])
def convert_to_docx():
    """Convert various formats TO DOCX"""
    
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part', 'error')
            return redirect(request.url)
        
        file = request.files['file']
        if file.filename == '':
            flash('No selected file', 'error')
            return redirect(request.url)
        
        if file:
            filename = secure_filename(file.filename)
            filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
            
            try:
                if not DOCX_AVAILABLE:
                    flash('DOCX conversion not supported. Please install python-docx library.', 'error')
                    return redirect(request.url)
                
                docx_buffer = convert_file_to_docx(filepath, filename, file_ext)
                
                if docx_buffer:
                    flash('Successfully converted to DOCX!', 'success')
                    return send_file(
                        docx_buffer, 
                        as_attachment=True, 
                        download_name=f"{filename.rsplit('.', 1)[0]}.docx"
                    )
                else:
                    flash(f'Unsupported file format: {file_ext}', 'error')
                    return redirect(request.url)
                    
            except Exception as e:
                current_app.logger.error(f"Error converting to DOCX: {e}")
                flash('Conversion failed. Please try again.', 'error')
                return redirect(request.url)
    
    return render_template('convert_to_docx.html')

def convert_file_to_pdf(filepath, filename, file_ext):
    """Convert various file types to PDF"""
    
    if file_ext == 'txt':
        return convert_txt_to_pdf(filepath)
    elif file_ext in ['md', 'markdown']:
        return convert_markdown_to_pdf(filepath)
    elif file_ext in ['html', 'htm']:
        return convert_html_to_pdf(filepath)
    elif file_ext == 'csv':
        return convert_csv_to_pdf(filepath)
    elif file_ext == 'json':
        return convert_json_to_pdf(filepath)
    elif file_ext in ['xlsx', 'xls']:
        return convert_excel_to_pdf(filepath)
    else:
        return None

def convert_file_to_docx(filepath, filename, file_ext):
    """Convert various file types to DOCX"""
    
    if file_ext == 'txt':
        return convert_txt_to_docx(filepath)
    elif file_ext in ['md', 'markdown']:
        return convert_markdown_to_docx(filepath)
    elif file_ext in ['html', 'htm']:
        return convert_html_to_docx(filepath)
    elif file_ext == 'csv':
        return convert_csv_to_docx(filepath)
    elif file_ext == 'json':
        return convert_json_to_docx(filepath)
    elif file_ext in ['xlsx', 'xls']:
        return convert_excel_to_docx(filepath)
    else:
        return None

# PDF Conversion Functions
def convert_txt_to_pdf(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    for para in content.split('\n\n'):
        if para.strip():
            p = Paragraph(para.replace('\n', '<br/>'), styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def convert_markdown_to_pdf(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Simple markdown parsing
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            header_text = line.lstrip('# ')
            style_name = f'Heading{min(level, 3)}'
            p = Paragraph(header_text, styles.get(style_name, styles['Heading1']))
            story.append(p)
        else:
            p = Paragraph(line, styles['Normal'])
            story.append(p)
        story.append(Spacer(1, 6))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def convert_html_to_pdf(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    if HTML_PARSER_AVAILABLE:
        soup = BeautifulSoup(content, 'html.parser')
        text_content = soup.get_text()
    else:
        text_content = content
    
    for para in text_content.split('\n\n'):
        if para.strip():
            p = Paragraph(para.replace('\n', '<br/>'), styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def convert_csv_to_pdf(filepath):
    df = pd.read_csv(filepath)
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    title = Paragraph("CSV Data", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))
    
    # Convert DataFrame to table data
    table_data = [df.columns.tolist()]
    table_data.extend(df.values.tolist())
    
    pdf_table = Table(table_data)
    pdf_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(pdf_table)
    doc.build(story)
    buffer.seek(0)
    return buffer

def convert_json_to_pdf(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        data = json.load(f)
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    title = Paragraph("JSON Data", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))
    
    json_text = json.dumps(data, indent=2)
    p = Paragraph(f"<pre>{json_text}</pre>", styles['Code'])
    story.append(p)
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def convert_excel_to_pdf(filepath):
    df = pd.read_excel(filepath)
    # Create a temporary CSV and use CSV conversion
    csv_path = filepath.replace('.xlsx', '.csv').replace('.xls', '.csv')
    df.to_csv(csv_path, index=False)
    result = convert_csv_to_pdf(csv_path)
    try:
        os.remove(csv_path)
    except:
        pass
    return result

# DOCX Conversion Functions
def convert_txt_to_docx(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    doc = Document()
    doc.add_heading('Text Document', 0)
    
    for para in content.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def convert_markdown_to_docx(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    doc = Document()
    
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            header_text = line.lstrip('# ')
            doc.add_heading(header_text, level)
        elif line.startswith('- ') or line.startswith('* '):
            list_text = line[2:]
            doc.add_paragraph(list_text, style='List Bullet')
        else:
            doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def convert_html_to_docx(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    doc = Document()
    doc.add_heading('HTML Document', 0)
    
    if HTML_PARSER_AVAILABLE:
        soup = BeautifulSoup(content, 'html.parser')
        text_content = soup.get_text()
    else:
        text_content = content
    
    for para in text_content.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def convert_csv_to_docx(filepath):
    df = pd.read_csv(filepath)
    
    doc = Document()
    doc.add_heading('CSV Data', 0)
    
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # Add header
    for i, column in enumerate(df.columns):
        table.cell(0, i).text = str(column)
        table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    # Add data
    for i, row in df.iterrows():
        for j, value in enumerate(row):
            table.cell(i + 1, j).text = str(value)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def convert_json_to_docx(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        data = json.load(f)
    
    doc = Document()
    doc.add_heading('JSON Data', 0)
    
    json_text = json.dumps(data, indent=2)
    doc.add_paragraph(json_text)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def convert_excel_to_docx(filepath):
    df = pd.read_excel(filepath)
    # Create a temporary CSV and use CSV conversion
    csv_path = filepath.replace('.xlsx', '.csv').replace('.xls', '.csv')
    df.to_csv(csv_path, index=False)
    result = convert_csv_to_docx(csv_path)
    try:
        os.remove(csv_path)
    except:
        pass
    return result
