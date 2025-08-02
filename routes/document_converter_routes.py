# routes/document_converter_routes.py
import os
import json
import re
from io import BytesIO
from flask import Blueprint, render_template, request, flash, redirect, send_file, current_app, jsonify
from docx import Document
from docx.shared import Inches
import pandas as pd
from werkzeug.utils import secure_filename
from utils.helpers import allowed_file

# Optional imports
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import markdown
    import markdown2
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    import lxml
    HTML_PARSER_AVAILABLE = True
except ImportError:
    HTML_PARSER_AVAILABLE = False

try:
    import yaml
    YAML_AVAILABLE = True
except ImportError:
    YAML_AVAILABLE = False

document_converter_bp = Blueprint('document_converter', __name__)

@document_converter_bp.route('/document-converter', methods=['GET', 'POST'])
def document_converter():
    from config import ALLOWED_CONVERTER_EXTENSIONS
    
    allowed_extensions = ALLOWED_CONVERTER_EXTENSIONS['doc']
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part', 'error')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            flash('No selected file', 'error')
            return redirect(request.url)

        if file and allowed_file(file.filename, allowed_extensions):
            filename = secure_filename(file.filename)
            filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            output_format = request.form.get('format')
            try:
                if output_format == 'pdf':
                    if not REPORTLAB_AVAILABLE:
                        flash('PDF conversion not supported. Please install reportlab library.', 'error')
                        return redirect(request.url)
                    
                    doc = Document(filepath)
                    pdf_buffer = BytesIO()
                    
                    # Create PDF using reportlab
                    c = canvas.Canvas(pdf_buffer, pagesize=letter)
                    width, height = letter
                    y_position = height - 50
                    
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text = para.text
                            # Simple text wrapping
                            lines = [text[i:i+80] for i in range(0, len(text), 80)]
                            for line in lines:
                                if y_position < 50:
                                    c.showPage()
                                    y_position = height - 50
                                c.drawString(50, y_position, line)
                                y_position -= 15
                    
                    c.save()
                    pdf_buffer.seek(0)
                    flash('Successfully converted to PDF!', 'success')
                    return send_file(pdf_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.pdf")
                elif output_format == 'csv':
                    result = convert_docx_to_csv(filepath)
                    csv_buffer = BytesIO(result.encode('utf-8'))
                    flash('Successfully converted to CSV!', 'success')
                    return send_file(csv_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.csv")
                elif output_format == 'html':
                    result = convert_docx_to_html(filepath)
                    html_buffer = BytesIO(result.encode('utf-8'))
                    flash('Successfully converted to HTML!', 'success')
                    return send_file(html_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.html")
                elif output_format == 'md':
                    result = convert_docx_to_markdown(filepath)
                    md_buffer = BytesIO(result.encode('utf-8'))
                    flash('Successfully converted to Markdown!', 'success')
                    return send_file(md_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.md")
                elif output_format == 'xml':
                    result = convert_docx_to_xml(filepath)
                    xml_buffer = BytesIO(result.encode('utf-8'))
                    flash('Successfully converted to XML!', 'success')
                    return send_file(xml_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.xml")
                elif output_format == 'yaml':
                    result = convert_docx_to_yaml(filepath)
                    yaml_buffer = BytesIO(result.encode('utf-8'))
                    flash('Successfully converted to YAML!', 'success')
                    return send_file(yaml_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.yaml")
                elif output_format == 'json':
                    result = extract_docx_data(filepath)
                    json_buffer = BytesIO(json.dumps(result, indent=2).encode('utf-8'))
                    flash('Successfully converted to JSON!', 'success')
                    return send_file(json_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.json")
                elif output_format == 'excel':
                    result = convert_docx_to_excel(filepath)
                    excel_buffer = BytesIO()
                    result.to_excel(excel_buffer, index=False, engine='openpyxl')
                    excel_buffer.seek(0)
                    flash('Successfully converted to Excel!', 'success')
                    return send_file(excel_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.xlsx")
                elif output_format == 'txt':
                    doc = Document(filepath)
                    full_text = [para.text for para in doc.paragraphs]
                    text = '\n'.join(full_text)
                    txt_buffer = BytesIO(text.encode('utf-8'))
                    flash('Successfully converted to TXT!', 'success')
                    return send_file(txt_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.txt")

            except Exception as e:
                flash(f"Error converting Document: {e}", 'error')
                return redirect(request.url)
        else:
            flash('Invalid file type. Please upload a DOCX file.', 'error')
            return redirect(request.url)
    
    return render_template('document_converter.html')

# Advanced Document Conversion Helper Functions

def convert_docx_to_csv(filepath):
    """Convert DOCX to CSV format by extracting paragraphs and tables"""
    doc = Document(filepath)
    data = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            data.append([para.text.strip()])
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):  # Only add non-empty rows
                data.append(row_data)
    
    # Convert to CSV format
    df = pd.DataFrame(data)
    return df.to_csv(index=False)

def convert_docx_to_html(filepath):
    """Convert DOCX to HTML with formatting"""
    doc = Document(filepath)
    html_content = """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Document Content</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }
        .paragraph { margin-bottom: 10px; }
        .bold { font-weight: bold; }
        .italic { font-style: italic; }
        table { border-collapse: collapse; width: 100%; margin: 20px 0; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        th { background-color: #f2f2f2; }
    </style>
</head>
<body>
"""
    
    # Convert paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # Basic style detection
            style_class = "paragraph"
            if para.style.name.startswith('Heading'):
                level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                html_content += f"<h{level}>{para.text}</h{level}>\n"
            else:
                html_content += f'<p class="{style_class}">{para.text}</p>\n'
    
    # Convert tables
    for table in doc.tables:
        html_content += "<table>\n"
        for i, row in enumerate(table.rows):
            if i == 0:
                html_content += "<thead><tr>"
                for cell in row.cells:
                    html_content += f"<th>{cell.text}</th>"
                html_content += "</tr></thead><tbody>\n"
            else:
                html_content += "<tr>"
                for cell in row.cells:
                    html_content += f"<td>{cell.text}</td>"
                html_content += "</tr>\n"
        html_content += "</tbody></table>\n"
    
    html_content += "</body>\n</html>"
    return html_content

def convert_docx_to_markdown(filepath):
    """Convert DOCX to Markdown format"""
    doc = Document(filepath)
    markdown_content = ""
    
    # Convert paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text = para.text.strip()
            
            # Handle different styles
            if para.style.name.startswith('Heading'):
                level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                markdown_content += "#" * int(level) + f" {text}\n\n"
            elif para.style.name == 'List Paragraph':
                markdown_content += f"- {text}\n"
            else:
                markdown_content += f"{text}\n\n"
    
    # Convert tables
    for table in doc.tables:
        if table.rows:
            # Header row
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            markdown_content += "| " + " | ".join(header_cells) + " |\n"
            markdown_content += "| " + " | ".join(["---"] * len(header_cells)) + " |\n"
            
            # Data rows
            for row in table.rows[1:]:
                row_cells = [cell.text.strip() for cell in row.cells]
                markdown_content += "| " + " | ".join(row_cells) + " |\n"
            
            markdown_content += "\n"
    
    return markdown_content

def convert_docx_to_xml(filepath):
    """Convert DOCX to structured XML"""
    doc = Document(filepath)
    xml_content = '<?xml version="1.0" encoding="UTF-8"?>\n<document>\n'
    
    # Document properties
    xml_content += '  <properties>\n'
    if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
        xml_content += f'    <title><![CDATA[{doc.core_properties.title}]]></title>\n'
    if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
        xml_content += f'    <author><![CDATA[{doc.core_properties.author}]]></author>\n'
    xml_content += '  </properties>\n'
    
    xml_content += '  <content>\n'
    
    # Convert paragraphs
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            style = para.style.name if para.style else 'Normal'
            xml_content += f'    <paragraph id="{i}" style="{style}">\n'
            xml_content += f'      <text><![CDATA[{para.text.strip()}]]></text>\n'
            xml_content += '    </paragraph>\n'
    
    # Convert tables
    for table_idx, table in enumerate(doc.tables):
        xml_content += f'    <table id="{table_idx}">\n'
        for row_idx, row in enumerate(table.rows):
            xml_content += f'      <row id="{row_idx}">\n'
            for cell_idx, cell in enumerate(row.cells):
                xml_content += f'        <cell id="{cell_idx}"><![CDATA[{cell.text.strip()}]]></cell>\n'
            xml_content += '      </row>\n'
        xml_content += '    </table>\n'
    
    xml_content += '  </content>\n</document>'
    return xml_content

def convert_docx_to_yaml(filepath):
    """Convert DOCX to YAML format"""
    if not YAML_AVAILABLE:
        return "# YAML library not available\n# Please install PyYAML to use this feature"
    
    doc = Document(filepath)
    data = {
        'document': {
            'properties': {},
            'paragraphs': [],
            'tables': []
        }
    }
    
    # Document properties
    if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
        data['document']['properties']['title'] = doc.core_properties.title
    if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
        data['document']['properties']['author'] = doc.core_properties.author
    
    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            para_data = {
                'style': para.style.name if para.style else 'Normal',
                'text': para.text.strip()
            }
            data['document']['paragraphs'].append(para_data)
    
    # Tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        data['document']['tables'].append(table_data)
    
    return yaml.dump(data, default_flow_style=False, allow_unicode=True)

def extract_docx_data(filepath):
    """Extract comprehensive data from DOCX"""
    doc = Document(filepath)
    data = {
        'metadata': {},
        'statistics': {},
        'content': {
            'paragraphs': [],
            'tables': [],
            'styles_used': []
        }
    }
    
    # Metadata
    core_props = doc.core_properties
    if core_props.title:
        data['metadata']['title'] = core_props.title
    if core_props.author:
        data['metadata']['author'] = core_props.author
    if core_props.subject:
        data['metadata']['subject'] = core_props.subject
    if core_props.created:
        data['metadata']['created'] = core_props.created.isoformat()
    if core_props.modified:
        data['metadata']['modified'] = core_props.modified.isoformat()
    
    # Statistics
    paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
    table_count = len(doc.tables)
    total_text = ' '.join([p.text for p in doc.paragraphs])
    word_count = len(total_text.split())
    
    data['statistics'] = {
        'paragraph_count': paragraph_count,
        'table_count': table_count,
        'word_count': word_count,
        'character_count': len(total_text)
    }
    
    # Content
    styles_used = set()
    for para in doc.paragraphs:
        if para.text.strip():
            style_name = para.style.name if para.style else 'Normal'
            styles_used.add(style_name)
            
            para_data = {
                'style': style_name,
                'text': para.text.strip(),
                'length': len(para.text.strip())
            }
            data['content']['paragraphs'].append(para_data)
    
    # Tables
    for table_idx, table in enumerate(doc.tables):
        table_data = {
            'table_id': table_idx,
            'rows': len(table.rows),
            'columns': len(table.columns) if table.rows else 0,
            'data': []
        }
        
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data['data'].append(row_data)
        
        data['content']['tables'].append(table_data)
    
    data['content']['styles_used'] = list(styles_used)
    
    return data

def convert_docx_to_excel(filepath):
    """Convert DOCX tables and content to Excel format"""
    doc = Document(filepath)
    
    # If there are tables, extract them
    if doc.tables:
        all_tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                df = pd.DataFrame(table_data[1:], columns=table_data[0] if table_data[0] else None)
                all_tables.append(df)
        
        if all_tables:
            return pd.concat(all_tables, ignore_index=True)
    
    # If no tables, convert paragraphs to Excel
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append({
                'Style': para.style.name if para.style else 'Normal',
                'Text': para.text.strip(),
                'Word_Count': len(para.text.split()),
                'Character_Count': len(para.text)
            })
    
    return pd.DataFrame(paragraphs)
