from flask import Blueprint, request, render_template, flash, send_file, current_app, jsonify, redirect, url_for
import os
import sys
import uuid
import tempfile
import threading
from werkzeug.utils import secure_filename
from datetime import datetime
import json

# Add project root to path for imports
project_root = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

# Configure Tesseract before importing pytesseract
try:
    from tesseract_config_helper import configure_tesseract
    configure_tesseract()
except Exception as e:
    print(f"Warning: Could not configure Tesseract helper: {e}")

# Try to import OCR dependencies, with fallbacks if not available
try:
    import pytesseract
    # Ensure Tesseract path is set
    if not pytesseract.pytesseract.tesseract_cmd or pytesseract.pytesseract.tesseract_cmd == 'tesseract':
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    
try:
    from PIL import Image, ImageEnhance, ImageFilter
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    
try:
    import cv2
    import numpy as np
    CV2_AVAILABLE = True
except ImportError:
    CV2_AVAILABLE = False
    
try:
    import fitz  # PyMuPDF for PDF processing
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    
try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False
    
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

text_ocr_bp = Blueprint('text_ocr', __name__)

ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'tiff', 'docx', 'txt'}
SUPPORTED_LANGUAGES = {
    'auto': 'Auto-detect',
    'eng': 'English',
    'hin': 'Hindi',
    'spa': 'Spanish',
    'fra': 'French',
    'deu': 'German',
    'ara': 'Arabic',
    'chi_sim': 'Chinese (Simplified)',
    'jpn': 'Japanese',
    'kor': 'Korean',
    'rus': 'Russian'
}

EXPORT_FORMATS = {
    'txt': {'extension': '.txt', 'mime_type': 'text/plain'},
    'docx': {'extension': '.docx', 'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'},
    'pdf': {'extension': '.pdf', 'mime_type': 'application/pdf'},
    'json': {'extension': '.json', 'mime_type': 'application/json'}
}

def allowed_file(filename):
    """Check if the uploaded file has an allowed extension"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def enhance_image(image, auto_enhance=True):
    """Enhance image quality for better OCR results using advanced techniques"""
    if not auto_enhance:
        return image
    
    try:
        # Convert PIL Image to numpy array
        img_array = np.array(image)
        
        # Convert to grayscale if not already
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        
        # Apply CLAHE (Contrast Limited Adaptive Histogram Equalization)
        # This improves contrast while preserving details
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)
        
        # Apply bilateral filter to denoise while preserving edges
        denoised = cv2.bilateralFilter(enhanced, 9, 75, 75)
        
        # Apply morphological operations to enhance text
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        morph = cv2.morphologyEx(denoised, cv2.MORPH_CLOSE, kernel)
        
        # Apply thresholding with adaptive method (Otsu's)
        # This works better than fixed threshold for varying lighting
        _, binary = cv2.threshold(morph, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Deskew the image (rotate if text is not horizontal)
        # Calculate the angle to deskew
        coords = np.column_stack(np.where(binary > 0))
        angle = cv2.minAreaRect(coords)[2]
        
        if angle < -45:
            angle = 90 + angle
        
        # Only rotate if angle is significant
        if abs(angle) > 0.5:
            h, w = binary.shape
            center = (w // 2, h // 2)
            rotation_matrix = cv2.getRotationMatrix2D(center, angle, 1.0)
            rotated = cv2.warpAffine(binary, rotation_matrix, (w, h), 
                                     borderMode=cv2.BORDER_REPLICATE)
        else:
            rotated = binary
        
        # Remove small noise (salt and pepper)
        cleaned = cv2.medianBlur(rotated, 3)
        
        # Apply invert if needed (ensure text is black on white)
        # Check if text is mostly dark
        if np.mean(cleaned) > 127:
            cleaned = cv2.bitwise_not(cleaned)
        
        # Dilate text slightly to make it bolder and easier to read
        kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (1, 1))
        enhanced_final = cv2.dilate(cleaned, kernel, iterations=1)
        
        # Convert back to PIL Image
        final_image = Image.fromarray(enhanced_final)
        
        return final_image
    except Exception as e:
        print(f"Image enhancement failed: {str(e)}")
        return image

def detect_language(text):
    """Detect the primary language of the extracted text"""
    try:
        from langdetect import detect
        lang_code = detect(text)
        
        # Map langdetect codes to tesseract codes
        lang_mapping = {
            'en': 'eng',
            'hi': 'hin',
            'es': 'spa',
            'fr': 'fra',
            'de': 'deu',
            'ar': 'ara',
            'zh-cn': 'chi_sim',
            'ja': 'jpn',
            'ko': 'kor',
            'ru': 'rus'
        }
        
        return lang_mapping.get(lang_code, 'eng')
    except:
        return 'eng'

def process_image_ocr(image_path, language='eng', auto_enhance=True):
    """Process image file for OCR"""
    if not TESSERACT_AVAILABLE:
        return {
            'error': '📷 Image OCR Not Available: OCR functionality is required for image files but Tesseract is not installed on your system. Please try uploading a PDF, DOCX, or TXT file instead - these work great without OCR! 📄',
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }
    
    if not PIL_AVAILABLE:
        return {
            'error': '📷 Image Processing Not Available: PIL/Pillow is required for image processing. Please install it and try again.',
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }
    
    try:
        # Open and enhance image
        image = Image.open(image_path)
        
        # Convert RGBA to RGB if necessary
        if image.mode == 'RGBA':
            rgb_image = Image.new('RGB', image.size, (255, 255, 255))
            rgb_image.paste(image, mask=image.split()[3])
            image = rgb_image
        
        # Enhance image if requested
        if auto_enhance and CV2_AVAILABLE:
            image = enhance_image(image, auto_enhance)
        
        # Perform OCR with optimized Tesseract configuration
        # These settings improve accuracy and speed:
        # --oem 1: Use new neural net OCR engine
        # --psm 1: Automatic page segmentation
        custom_config = r'--oem 1 --psm 1'
        
        # Set a timeout for OCR to prevent hanging
        ocr_result = {'text': '', 'error': None, 'language': language}
        
        def run_ocr():
            nonlocal language
            try:
                if language != 'auto':
                    ocr_result['text'] = pytesseract.image_to_string(image, lang=language, config=custom_config)
                else:
                    # Try to detect language from a sample
                    sample_text = pytesseract.image_to_string(image, lang='eng', config=custom_config)[:500]
                    detected_lang = detect_language(sample_text)
                    ocr_result['text'] = pytesseract.image_to_string(image, lang=detected_lang, config=custom_config)
                    # Store detected language
                    language = detected_lang
                    ocr_result['language'] = detected_lang
            except Exception as e:
                ocr_result['error'] = str(e)
        
        # Run OCR in a thread with timeout (30 seconds)
        ocr_thread = threading.Thread(target=run_ocr, daemon=True)
        ocr_thread.start()
        ocr_thread.join(timeout=30)
        
        if ocr_thread.is_alive():
            return {
                'error': 'OCR processing timeout: The image took too long to process. Please try a simpler image or smaller file.',
                'text': '',
                'confidence': 0,
                'language': 'Unknown',
                'word_count': 0
            }
        
        if ocr_result['error']:
            raise Exception(ocr_result['error'])
        
        text = ocr_result['text']
        
        # Default confidence based on text quality
        avg_confidence = 85  # Default confidence for OCR
        if len(text.strip()) < 10:
            avg_confidence = 60  # Lower confidence for very short extractions
        
        # Clean up the extracted text
        # Remove extra whitespace and improve formatting
        text = text.strip()
        # Replace multiple spaces with single space
        text = ' '.join(text.split())
        
        return {
            'text': text,
            'confidence': round(avg_confidence, 2),
            'language': SUPPORTED_LANGUAGES.get(language, language),
            'word_count': len(text.split())
        }
        
    except pytesseract.TesseractNotFoundError:
        return {
            'error': '📷 Tesseract Not Found: Tesseract OCR is not installed or not in your PATH. Image OCR is not available. Please try uploading a PDF, DOCX, or TXT file instead! 📄',
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }
    except Exception as e:
        return {
            'error': f"OCR processing failed: {str(e)}",
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }

def process_pdf_ocr(pdf_path, language='eng', auto_enhance=True):
    """Process PDF file for OCR"""
    if not PYMUPDF_AVAILABLE:
        return {
            'error': 'PDF processing dependencies not available. Please install PyMuPDF.',
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }
    
    try:
        all_text = []
        pdf_document = fitz.open(pdf_path)
        total_confidence = 0
        page_count = 0
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # First try to extract text directly
            text = page.get_text()
            
            if text.strip():
                # If text exists, it's already a text PDF
                all_text.append(text)
                total_confidence += 100  # High confidence for native text
            else:
                # Convert page to image for OCR (with higher quality)
                mat = fitz.Matrix(3, 3)  # Increase resolution to 3x for better OCR
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_data = pix.pil_tobytes(format="PNG")
                
                # Create temporary image file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                    tmp_img.write(img_data)
                    tmp_img_path = tmp_img.name
                
                # Process image with OCR
                result = process_image_ocr(tmp_img_path, language, auto_enhance)
                all_text.append(result['text'])
                total_confidence += result['confidence']
                
                # Clean up temporary file
                try:
                    os.unlink(tmp_img_path)
                except:
                    pass
            
            page_count += 1
        
        pdf_document.close()
        
        full_text = '\n\n'.join(all_text)
        # Clean up text
        full_text = full_text.strip()
        # Remove excessive whitespace while preserving paragraphs
        lines = [line.strip() for line in full_text.split('\n') if line.strip()]
        full_text = '\n'.join(lines)
        
        avg_confidence = total_confidence / page_count if page_count > 0 else 0
        
        return {
            'text': full_text,
            'confidence': round(avg_confidence, 2),
            'language': SUPPORTED_LANGUAGES.get(language, language),
            'word_count': len(full_text.split())
        }
        
    except Exception as e:
        return {
            'error': f"PDF processing failed: {str(e)}",
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }

def process_docx_ocr(docx_path):
    """Process DOCX file for text extraction"""
    if not DOCX2TXT_AVAILABLE:
        return {
            'error': 'DOCX processing dependencies not available. Please install docx2txt.',
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }
    
    try:
        text = docx2txt.process(docx_path)
        
        # Clean up the extracted text
        text = text.strip()
        # Remove excessive whitespace while preserving paragraphs
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        text = '\n'.join(lines)
        
        return {
            'text': text,
            'confidence': 100,  # Direct text extraction, high confidence
            'language': detect_language(text[:500]) if text else 'eng',
            'word_count': len(text.split())
        }
    except Exception as e:
        return {
            'error': f"DOCX processing failed: {str(e)}",
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }

def export_to_txt(text, output_path):
    """Export text to TXT format"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    return True

def export_to_docx(text, output_path):
    """Export text to DOCX format"""
    if not PYTHON_DOCX_AVAILABLE:
        return False
        
    try:
        doc = Document()
        
        # Add title
        doc.add_heading('OCR Extracted Text', 0)
        
        # Add metadata
        doc.add_paragraph(f'Extraction Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')
        
        # Add extracted text
        for paragraph in text.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"DOCX export failed: {str(e)}")
        return False

def export_to_pdf(text, output_path):
    """Export text to searchable PDF format"""
    if not REPORTLAB_AVAILABLE:
        return False
    
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.units import inch
        import html
        
        # Create a PDF document
        doc = SimpleDocTemplate(output_path, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title = Paragraph("OCR Extracted Text", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 0.2*inch))
        
        # Add extraction date
        date_para = Paragraph(f"Extraction Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal'])
        story.append(date_para)
        story.append(Spacer(1, 0.2*inch))
        
        # Add text content - properly escape HTML special characters
        paragraphs = text.split('\n\n')
        for idx, paragraph in enumerate(paragraphs):
            if paragraph.strip():
                # Escape HTML special characters to prevent parsing errors
                clean_text = html.escape(paragraph.strip())
                # Replace newlines with <br/> tags properly
                clean_text = clean_text.replace('\n', '<br/>')
                
                try:
                    para = Paragraph(clean_text, styles['Normal'])
                    story.append(para)
                    story.append(Spacer(1, 0.1*inch))
                    
                    # Add page break every 10 paragraphs to avoid overflow
                    if (idx + 1) % 10 == 0:
                        story.append(PageBreak())
                except Exception as para_error:
                    print(f"Skipping problematic paragraph: {para_error}")
                    # Add as plain text if Paragraph fails
                    story.append(Paragraph(html.escape(paragraph.strip()[:200]), styles['Normal']))
        
        # Build the PDF
        doc.build(story)
        return True
    except Exception as e:
        print(f"PDF export failed: {str(e)}")
        return False

def export_to_json(ocr_result, output_path):
    """Export OCR result to JSON format"""
    try:
        json_data = {
            'metadata': {
                'extraction_date': datetime.now().isoformat(),
                'confidence': ocr_result['confidence'],
                'language': ocr_result['language'],
                'word_count': ocr_result['word_count']
            },
            'text': ocr_result['text'],
            'paragraphs': ocr_result['text'].split('\n\n')
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        
        return True
    except Exception as e:
        print(f"JSON export failed: {str(e)}")
        return False

@text_ocr_bp.route('/convert/text-ocr')
@text_ocr_bp.route('/convert/text-ocr/')
def text_ocr():
    """Display the text & OCR processor page"""
    return render_template('text_ocr_converters/text_ocr.html')

@text_ocr_bp.route('/convert/text-ocr', methods=['POST'])
@text_ocr_bp.route('/convert/text-ocr/', methods=['POST'])
def process_ocr():
    """Handle OCR processing"""
    try:
        # Check if file was uploaded
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded', 'success': False}), 400
        
        file = request.files['file']
        language = request.form.get('ocr_language', 'auto')
        output_format = request.form.get('output_format', 'txt')
        auto_enhance = request.form.get('remove_noise') == 'on'
        
        # Validate file
        if file.filename == '':
            return jsonify({'error': 'No file selected', 'success': False}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Please upload PDF, PNG, JPG, JPEG, TIFF, DOCX, or TXT files', 'success': False}), 400
        
        if output_format not in EXPORT_FORMATS:
            return jsonify({'error': 'Invalid output format selected', 'success': False}), 400
        
        # Get file extension
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        # Check if it's an image file and OCR is not available
        if file_extension in ['.png', '.jpg', '.jpeg', '.tiff']:
            if not TESSERACT_AVAILABLE:
                return jsonify({
                    'error': '📷 Image OCR Not Available: Tesseract OCR is not installed on your system. Please try uploading a PDF, DOCX, or TXT file instead - these work great without OCR! 📄',
                    'success': False
                }), 400
        
        # Generate unique filename
        unique_id = str(uuid.uuid4())
        filename = secure_filename(file.filename)
        base_name = os.path.splitext(filename)[0]
        
        # Save uploaded file
        upload_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
        file.save(upload_path)
        
        # Process file based on type
        if file_extension in ['.png', '.jpg', '.jpeg', '.tiff']:
            ocr_result = process_image_ocr(upload_path, language, auto_enhance)
        elif file_extension == '.pdf':
            ocr_result = process_pdf_ocr(upload_path, language, auto_enhance)
        elif file_extension == '.docx':
            ocr_result = process_docx_ocr(upload_path)
        elif file_extension == '.txt':
            # For TXT files, just read the content
            with open(upload_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            ocr_result = {
                'text': text_content,
                'confidence': 100,
                'language': detect_language(text_content[:500]) if text_content else 'Unknown',
                'word_count': len(text_content.split())
            }
        else:
            ocr_result = {'error': 'Unsupported file type'}
        
        # Clean up uploaded file
        if os.path.exists(upload_path):
            os.remove(upload_path)
        
        # Check for errors
        if 'error' in ocr_result:
            return jsonify({'error': ocr_result['error'], 'success': False}), 500
        
        # Export to requested format
        output_extension = EXPORT_FORMATS[output_format]['extension']
        output_filename = f"{base_name}_ocr_{unique_id}{output_extension}"
        output_path = os.path.join(current_app.config['OUTPUT_FOLDER'], output_filename)
        
        export_success = False
        if output_format == 'txt':
            export_success = export_to_txt(ocr_result['text'], output_path)
        elif output_format == 'docx':
            export_success = export_to_docx(ocr_result['text'], output_path)
        elif output_format == 'pdf':
            export_success = export_to_pdf(ocr_result['text'], output_path)
        elif output_format == 'json':
            export_success = export_to_json(ocr_result, output_path)
        
        if not export_success:
            return jsonify({'error': 'Export failed. Please try again.', 'success': False}), 500
        
        # Return success response with download URL
        return jsonify({
            'success': True,
            'message': f'OCR processing successful! {ocr_result["word_count"]} words extracted.',
            'download_url': f'/ocr-download/{output_filename}',
            'ocr_result': {
                'word_count': ocr_result['word_count'],
                'confidence': ocr_result['confidence'],
                'language': ocr_result['language']
            }
        })
        
    except Exception as e:
        return jsonify({'error': f'An error occurred: {str(e)}', 'success': False}), 500

@text_ocr_bp.route('/ocr-download/<filename>')
def download_file(filename):
    """Download processed file"""
    try:
        file_path = os.path.join(current_app.config['OUTPUT_FOLDER'], filename)
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found', 'success': False}), 404
        
        # Determine mime type from filename
        file_extension = os.path.splitext(filename)[1].lower()
        mime_type = 'application/octet-stream'  # default
        
        for format_info in EXPORT_FORMATS.values():
            if format_info['extension'] == file_extension:
                mime_type = format_info['mime_type']
                break
        
        return send_file(file_path, 
                        as_attachment=True,
                        download_name=filename,
                        mimetype=mime_type)
        
    except Exception as e:
        return jsonify({'error': f'Download failed: {str(e)}', 'success': False}), 500

@text_ocr_bp.route('/api/ocr', methods=['POST'])
def api_ocr():
    """API endpoint for OCR processing"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded', 'success': False}), 400
        
        file = request.files['file']
        language = request.form.get('ocr_language', 'auto')
        output_format = request.form.get('output_format', 'json')
        auto_enhance = request.form.get('remove_noise', 'true').lower() == 'true'
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type', 'success': False}), 400
        
        # Process file
        unique_id = str(uuid.uuid4())
        filename = secure_filename(file.filename)
        upload_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
        file.save(upload_path)
        
        # Get file extension
        file_extension = os.path.splitext(filename)[1].lower()
        
        # Process based on file type
        if file_extension in ['.png', '.jpg', '.jpeg', '.tiff']:
            ocr_result = process_image_ocr(upload_path, language, auto_enhance)
        elif file_extension == '.pdf':
            ocr_result = process_pdf_ocr(upload_path, language, auto_enhance)
        elif file_extension == '.docx':
            ocr_result = process_docx_ocr(upload_path)
        else:
            ocr_result = {'error': 'Unsupported file type'}
        
        os.remove(upload_path)
        
        if 'error' in ocr_result:
            return jsonify({'error': ocr_result['error'], 'success': False}), 500
        
        # Return JSON response
        return jsonify({
            'success': True,
            'text': ocr_result['text'],
            'confidence': ocr_result['confidence'],
            'language': ocr_result['language'],
            'word_count': ocr_result['word_count']
        })
        
    except Exception as e:
        return jsonify({'error': str(e), 'success': False}), 500
