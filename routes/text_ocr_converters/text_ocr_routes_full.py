from flask import Blueprint, request, render_template, flash, send_file, current_app, jsonify, redirect, url_for
import os
import uuid
import tempfile
from werkzeug.utils import secure_filename
from datetime import datetime
import json

# Try to import OCR dependencies, with fallbacks if not available
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    
try:
    from PIL import Image, ImageEnhance, ImageFilter
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    
try:
    import cv2
    import numpy as np
    CV2_AVAILABLE = True
except ImportError:
    CV2_AVAILABLE = False
    
try:
    import fitz  # PyMuPDF for PDF processing
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    
try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False
    
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

text_ocr_bp = Blueprint('text_ocr', __name__)

ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'tiff', 'docx', 'txt'}
SUPPORTED_LANGUAGES = {
    'auto': 'Auto-detect',
    'eng': 'English',
    'hin': 'Hindi',
    'spa': 'Spanish',
    'fra': 'French',
    'deu': 'German',
    'ara': 'Arabic',
    'chi_sim': 'Chinese (Simplified)',
    'jpn': 'Japanese',
    'kor': 'Korean',
    'rus': 'Russian'
}

EXPORT_FORMATS = {
    'txt': {'extension': '.txt', 'mime_type': 'text/plain'},
    'docx': {'extension': '.docx', 'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'},
    'pdf': {'extension': '.pdf', 'mime_type': 'application/pdf'},
    'json': {'extension': '.json', 'mime_type': 'application/json'}
}

def allowed_file(filename):
    """Check if the uploaded file has an allowed extension"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def enhance_image(image, auto_enhance=True):
    """Enhance image quality for better OCR results"""
    if not auto_enhance:
        return image
    
    try:
        # Convert PIL Image to numpy array
        img_array = np.array(image)
        
        # Convert to grayscale if not already
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        
        # Apply thresholding to get better contrast
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Noise removal
        denoised = cv2.medianBlur(thresh, 3)
        
        # Convert back to PIL Image
        enhanced_image = Image.fromarray(denoised)
        
        return enhanced_image
    except Exception as e:
        print(f"Image enhancement failed: {str(e)}")
        return image

def detect_language(text):
    """Detect the primary language of the extracted text"""
    try:
        from langdetect import detect
        lang_code = detect(text)
        
        # Map langdetect codes to tesseract codes
        lang_mapping = {
            'en': 'eng',
            'hi': 'hin',
            'es': 'spa',
            'fr': 'fra',
            'de': 'deu',
            'ar': 'ara',
            'zh-cn': 'chi_sim',
            'ja': 'jpn',
            'ko': 'kor',
            'ru': 'rus'
        }
        
        return lang_mapping.get(lang_code, 'eng')
    except:
        return 'eng'

def process_image_ocr(image_path, language='eng', auto_enhance=True):
    """Process image file for OCR"""
    if not TESSERACT_AVAILABLE or not PIL_AVAILABLE:
        return {
            'error': 'OCR dependencies not available. Please install pytesseract and PIL.',
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }
    
    try:
        # Open and enhance image
        image = Image.open(image_path)
        
        # Convert RGBA to RGB if necessary
        if image.mode == 'RGBA':
            rgb_image = Image.new('RGB', image.size, (255, 255, 255))
            rgb_image.paste(image, mask=image.split()[3])
            image = rgb_image
        
        # Enhance image if requested
        if auto_enhance and CV2_AVAILABLE:
            image = enhance_image(image, auto_enhance)
        
        # Perform OCR
        custom_config = r'--oem 3 --psm 6'
        if language != 'auto':
            text = pytesseract.image_to_string(image, lang=language, config=custom_config)
        else:
            # Try to detect language from a sample
            sample_text = pytesseract.image_to_string(image, lang='eng', config=custom_config)[:500]
            detected_lang = detect_language(sample_text)
            text = pytesseract.image_to_string(image, lang=detected_lang, config=custom_config)
            language = detected_lang
        
        # Get confidence data
        try:
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 75
        except:
            avg_confidence = 75  # Default confidence if detection fails
        
        return {
            'text': text.strip(),
            'confidence': round(avg_confidence, 2),
            'language': SUPPORTED_LANGUAGES.get(language, language),
            'word_count': len(text.split())
        }
        
    except Exception as e:
        return {
            'error': f"OCR processing failed: {str(e)}",
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }

def process_pdf_ocr(pdf_path, language='eng', auto_enhance=True):
    """Process PDF file for OCR"""
    if not PYMUPDF_AVAILABLE:
        return {
            'error': 'PDF processing dependencies not available. Please install PyMuPDF.',
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }
    
    try:
        all_text = []
        pdf_document = fitz.open(pdf_path)
        total_confidence = 0
        page_count = 0
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # First try to extract text directly
            text = page.get_text()
            
            if text.strip():
                # If text exists, it's already a text PDF
                all_text.append(text)
                total_confidence += 100  # High confidence for native text
            else:
                # Convert page to image for OCR
                mat = fitz.Matrix(2, 2)  # Increase resolution
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.pil_tobytes(format="PNG")
                
                # Create temporary image file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                    tmp_img.write(img_data)
                    tmp_img_path = tmp_img.name
                
                # Process image with OCR
                result = process_image_ocr(tmp_img_path, language, auto_enhance)
                all_text.append(result['text'])
                total_confidence += result['confidence']
                
                # Clean up temporary file
                os.unlink(tmp_img_path)
            
            page_count += 1
        
        pdf_document.close()
        
        full_text = '\n\n'.join(all_text)
        avg_confidence = total_confidence / page_count if page_count > 0 else 0
        
        return {
            'text': full_text.strip(),
            'confidence': round(avg_confidence, 2),
            'language': SUPPORTED_LANGUAGES.get(language, language),
            'word_count': len(full_text.split())
        }
        
    except Exception as e:
        return {
            'error': f"PDF processing failed: {str(e)}",
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }

def process_docx_ocr(docx_path):
    """Process DOCX file for text extraction"""
    if not DOCX2TXT_AVAILABLE:
        return {
            'error': 'DOCX processing dependencies not available. Please install docx2txt.',
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }
    
    try:
        text = docx2txt.process(docx_path)
        
        return {
            'text': text.strip(),
            'confidence': 100,  # Direct text extraction, high confidence
            'language': detect_language(text[:500]) if text else 'eng',
            'word_count': len(text.split())
        }
    except Exception as e:
        return {
            'error': f"DOCX processing failed: {str(e)}",
            'text': '',
            'confidence': 0,
            'language': 'Unknown',
            'word_count': 0
        }

def export_to_txt(text, output_path):
    """Export text to TXT format"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    return True

def export_to_docx(text, output_path):
    """Export text to DOCX format"""
    if not PYTHON_DOCX_AVAILABLE:
        return False
        
    try:
        doc = Document()
        
        # Add title
        doc.add_heading('OCR Extracted Text', 0)
        
        # Add metadata
        doc.add_paragraph(f'Extraction Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')
        
        # Add extracted text
        for paragraph in text.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"DOCX export failed: {str(e)}")
        return False

def export_to_pdf(text, output_path):
    """Export text to searchable PDF format"""
    if not REPORTLAB_AVAILABLE:
        return False
    
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        
        # Create a PDF document
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title = Paragraph("OCR Extracted Text", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 0.2*inch))
        
        # Add extraction date
        date_para = Paragraph(f"Extraction Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal'])
        story.append(date_para)
        story.append(Spacer(1, 0.2*inch))
        
        # Add text content
        for paragraph in text.split('\n\n'):
            if paragraph.strip():
                para = Paragraph(paragraph.replace('\n', '<br/>'), styles['Normal'])
                story.append(para)
                story.append(Spacer(1, 0.1*inch))
        
        # Build the PDF
        doc.build(story)
        return True
    except Exception as e:
        print(f"PDF export failed: {str(e)}")
        return False

def export_to_json(ocr_result, output_path):
    """Export OCR result to JSON format"""
    try:
        json_data = {
            'metadata': {
                'extraction_date': datetime.now().isoformat(),
                'confidence': ocr_result['confidence'],
                'language': ocr_result['language'],
                'word_count': ocr_result['word_count']
            },
            'text': ocr_result['text'],
            'paragraphs': ocr_result['text'].split('\n\n')
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        
        return True
    except Exception as e:
        print(f"JSON export failed: {str(e)}")
        return False

@text_ocr_bp.route('/text-ocr')
def text_ocr():
    """Display the text & OCR processor page"""
    return render_template('text_ocr_converters/text_ocr.html')

@text_ocr_bp.route('/text-ocr', methods=['POST'])
def process_ocr():
    """Handle OCR processing"""
    try:
        # Check if file was uploaded
        if 'ocr_file' not in request.files:
            flash('No file uploaded', 'error')
            return redirect(url_for('text_ocr.text_ocr'))
        
        file = request.files['ocr_file']
        language = request.form.get('language', 'auto')
        output_format = request.form.get('output_format', 'txt')
        auto_enhance = request.form.get('auto_enhance') == 'on'
        
        # Validate file
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(url_for('text_ocr.text_ocr'))
        
        if not allowed_file(file.filename):
            flash('Invalid file type. Please upload PDF, PNG, JPG, JPEG, TIFF, DOCX, or TXT files', 'error')
            return redirect(url_for('text_ocr.text_ocr'))
        
        if output_format not in EXPORT_FORMATS:
            flash('Invalid output format selected', 'error')
            return redirect(url_for('text_ocr.text_ocr'))
        
        # Generate unique filename
        unique_id = str(uuid.uuid4())
        filename = secure_filename(file.filename)
        base_name = os.path.splitext(filename)[0]
        file_extension = os.path.splitext(filename)[1].lower()
        
        # Save uploaded file
        upload_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
        file.save(upload_path)
        
        # Process file based on type
        if file_extension in ['.png', '.jpg', '.jpeg', '.tiff']:
            ocr_result = process_image_ocr(upload_path, language, auto_enhance)
        elif file_extension == '.pdf':
            ocr_result = process_pdf_ocr(upload_path, language, auto_enhance)
        elif file_extension == '.docx':
            ocr_result = process_docx_ocr(upload_path)
        elif file_extension == '.txt':
            # For TXT files, just read the content
            with open(upload_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            ocr_result = {
                'text': text_content,
                'confidence': 100,
                'language': detect_language(text_content[:500]) if text_content else 'Unknown',
                'word_count': len(text_content.split())
            }
        else:
            ocr_result = {'error': 'Unsupported file type'}
        
        # Clean up uploaded file
        os.remove(upload_path)
        
        # Check for errors
        if 'error' in ocr_result:
            flash(ocr_result['error'], 'error')
            return redirect(url_for('text_ocr.text_ocr'))
        
        # Export to requested format
        output_extension = EXPORT_FORMATS[output_format]['extension']
        output_filename = f"{base_name}_ocr_{unique_id}{output_extension}"
        output_path = os.path.join(current_app.config['OUTPUT_FOLDER'], output_filename)
        
        export_success = False
        if output_format == 'txt':
            export_success = export_to_txt(ocr_result['text'], output_path)
        elif output_format == 'docx':
            export_success = export_to_docx(ocr_result['text'], output_path)
        elif output_format == 'pdf':
            export_success = export_to_pdf(ocr_result['text'], output_path)
        elif output_format == 'json':
            export_success = export_to_json(ocr_result, output_path)
        
        if not export_success:
            flash('Export failed. Please try again.', 'error')
            return redirect(url_for('text_ocr.text_ocr'))
        
        # Return success response with OCR results
        flash(f'OCR processing successful! {ocr_result["word_count"]} words extracted.', 'success')
        return render_template('text_ocr_converters/text_ocr.html', 
                             ocr_result=ocr_result,
                             download_file=output_filename,
                             output_format=output_format)
        
    except Exception as e:
        flash(f'An error occurred: {str(e)}', 'error')
        return redirect(url_for('text_ocr.text_ocr'))

@text_ocr_bp.route('/ocr-download/<filename>')
def download_file(filename):
    """Download processed file"""
    try:
        file_path = os.path.join(current_app.config['OUTPUT_FOLDER'], filename)
        
        if not os.path.exists(file_path):
            flash('File not found', 'error')
            return redirect(url_for('text_ocr.text_ocr'))
        
        # Determine mime type from filename
        file_extension = os.path.splitext(filename)[1].lower()
        mime_type = 'application/octet-stream'  # default
        
        for format_info in EXPORT_FORMATS.values():
            if format_info['extension'] == file_extension:
                mime_type = format_info['mime_type']
                break
        
        return send_file(file_path, 
                        as_attachment=True,
                        download_name=filename,
                        mimetype=mime_type)
        
    except Exception as e:
        flash(f'Download failed: {str(e)}', 'error')
        return redirect(url_for('text_ocr.text_ocr'))

@text_ocr_bp.route('/api/ocr', methods=['POST'])
def api_ocr():
    """API endpoint for OCR processing"""
    try:
        if 'ocr_file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['ocr_file']
        language = request.form.get('language', 'auto')
        output_format = request.form.get('output_format', 'json')
        auto_enhance = request.form.get('auto_enhance', 'true').lower() == 'true'
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type'}), 400
        
        # Process file
        unique_id = str(uuid.uuid4())
        filename = secure_filename(file.filename)
        upload_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"{unique_id}_{filename}")
        file.save(upload_path)
        
        # Get file extension
        file_extension = os.path.splitext(filename)[1].lower()
        
        # Process based on file type
        if file_extension in ['.png', '.jpg', '.jpeg', '.tiff']:
            ocr_result = process_image_ocr(upload_path, language, auto_enhance)
        elif file_extension == '.pdf':
            ocr_result = process_pdf_ocr(upload_path, language, auto_enhance)
        elif file_extension == '.docx':
            ocr_result = process_docx_ocr(upload_path)
        else:
            ocr_result = {'error': 'Unsupported file type'}
        
        os.remove(upload_path)
        
        if 'error' in ocr_result:
            return jsonify({'error': ocr_result['error']}), 500
        
        # Return JSON response
        return jsonify({
            'success': True,
            'text': ocr_result['text'],
            'confidence': ocr_result['confidence'],
            'language': ocr_result['language'],
            'word_count': ocr_result['word_count']
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500