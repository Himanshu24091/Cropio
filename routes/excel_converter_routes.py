# routes/excel_converter_routes.py
import os
from io import BytesIO
from flask import Blueprint, render_template, request, flash, redirect, send_file, current_app, jsonify
import pandas as pd
from werkzeug.utils import secure_filename
from utils.helpers import allowed_file
import json
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

excel_converter_bp = Blueprint('excel_converter', __name__)

def convert_excel_to_pdf(df, filename):
    """Convert DataFrame to PDF using ReportLab"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    
    # Create story list to hold flowable objects
    story = []
    
    # Add title
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    
    title = Paragraph(f"Excel Data: {filename.rsplit('.', 1)[0]}", title_style)
    story.append(title)
    story.append(Spacer(1, 12))
    
    # Convert DataFrame to table data
    data = [df.columns.tolist()]  # Header
    data.extend(df.values.tolist())  # Data rows
    
    # Create table
    table = Table(data)
    
    # Style the table
    table_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ])
    
    table.setStyle(table_style)
    story.append(table)
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def convert_excel_to_docx(df, filename):
    """Convert DataFrame to DOCX using python-docx"""
    buffer = BytesIO()
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Excel Data: {filename.rsplit(".", 1)[0]}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some space
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)
        # Make header bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value) if pd.notna(value) else ''
    
    # Adjust column widths
    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(1.5)
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@excel_converter_bp.route('/excel-converter', methods=['GET', 'POST'])
def excel_converter():
    from config import ALLOWED_CONVERTER_EXTENSIONS
    
    allowed_extensions = ALLOWED_CONVERTER_EXTENSIONS['excel']
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part', 'error')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            flash('No selected file', 'error')
            return redirect(request.url)

        if file and allowed_file(file.filename, allowed_extensions):
            filename = secure_filename(file.filename)
            filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            output_format = request.form.get('format')
            try:
                # Try to read Excel file with multiple sheets support
                excel_file = pd.ExcelFile(filepath)
                sheet_name = request.form.get('sheet_name', excel_file.sheet_names[0] if excel_file.sheet_names else 0)
                df = pd.read_excel(filepath, sheet_name=sheet_name)
                if output_format == 'csv':
                    csv_buffer = BytesIO()
                    df.to_csv(csv_buffer, index=False)
                    csv_buffer.seek(0)
                    flash('Successfully converted to CSV!', 'success')
                    return send_file(csv_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.csv")
                elif output_format == 'json':
                    json_buffer = BytesIO(df.to_json(orient='records').encode('utf-8'))
                    flash('Successfully converted to JSON!', 'success')
                    return send_file(json_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.json")
                elif output_format == 'html':
                    html_buffer = BytesIO(df.to_html().encode('utf-8'))
                    flash('Successfully converted to HTML!', 'success')
                    return send_file(html_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.html")
                elif output_format == 'md':
                    md_text = df.to_markdown()
                    md_buffer = BytesIO(md_text.encode('utf-8'))
                    flash('Successfully converted to Markdown!', 'success')
                    return send_file(md_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.md")
                elif output_format == 'xml':
                    xml_text = df.to_xml()
                    xml_buffer = BytesIO(xml_text.encode('utf-8'))
                    flash('Successfully converted to XML!', 'success')
                    return send_file(xml_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.xml")
                elif output_format == 'latex':
                    latex_text = df.to_latex()
                    latex_buffer = BytesIO(latex_text.encode('utf-8'))
                    flash('Successfully converted to LaTeX!', 'success')
                    return send_file(latex_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.tex")
                elif output_format == 'pdf':
                    pdf_buffer = convert_excel_to_pdf(df, filename)
                    flash('Successfully converted to PDF!', 'success')
                    return send_file(pdf_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.pdf")
                elif output_format == 'docx':
                    docx_buffer = convert_excel_to_docx(df, filename)
                    flash('Successfully converted to DOCX!', 'success')
                    return send_file(docx_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}.docx")
            except Exception as e:
                flash(f"Error converting Excel: {e}", 'error')
                return redirect(request.url)
        else:
            flash('Invalid file type. Please upload an Excel file (XLSX, XLS).', 'error')
            return redirect(request.url)
            
    return render_template('excel_converter.html')
