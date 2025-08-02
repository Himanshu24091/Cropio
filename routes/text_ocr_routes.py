# routes/text_ocr_routes.py
import os
from io import BytesIO
from flask import Blueprint, render_template, request, flash, redirect, send_file, current_app, jsonify
import pytesseract
from PIL import Image
from werkzeug.utils import secure_filename
from utils.helpers import allowed_file
import json
import docx
from docx import Document
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

text_ocr_bp = Blueprint('text_ocr', __name__)

# Set Tesseract path (adjust path based on your installation)
# For Windows, uncomment and adjust the path below:
import platform
if platform.system() == 'Windows':
    # Try common installation paths
    import os
    possible_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        r'C:\Users\%s\AppData\Local\Tesseract-OCR\tesseract.exe' % os.getenv('USERNAME')
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            break
    else:
        # If not found in common paths, try to find it in PATH
        import shutil
        tesseract_path = shutil.which('tesseract')
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path

def check_tesseract_installation():
    """Check if Tesseract is properly installed and configured"""
    try:
        # Try to get Tesseract version
        version = pytesseract.get_tesseract_version()
        return True, f"Tesseract version: {version}"
    except Exception as e:
        return False, f"Tesseract not found: {str(e)}"

def extract_text_from_image(image_path):
    """Extract text from image using OCR"""
    try:
        # First check if Tesseract is available
        is_available, message = check_tesseract_installation()
        if not is_available:
            raise Exception(f"Tesseract is not installed or not in your PATH. {message}. See TESSERACT_SETUP.md file for more information.")
        
        image = Image.open(image_path)
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Extract text using pytesseract
        extracted_text = pytesseract.image_to_string(image)
        return extracted_text.strip()
    except Exception as e:
        if "tesseract is not installed" in str(e).lower() or "not in your path" in str(e).lower():
            raise Exception(f"tesseract is not installed or it's not in your PATH. See README file for more information.")
        else:
            raise Exception(f"OCR processing failed: {str(e)}")

def create_text_pdf(text, filename):
    """Create PDF from text"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    
    # Create story list
    story = []
    
    # Add title
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    
    title = Paragraph(f"Extracted Text: {filename.rsplit('.', 1)[0]}", title_style)
    story.append(title)
    story.append(Spacer(1, 12))
    
    # Add text content
    content_style = ParagraphStyle(
        'Content',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=12,
        leftIndent=20,
        rightIndent=20
    )
    
    # Split text into paragraphs
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para.strip(), content_style))
            story.append(Spacer(1, 6))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_text_docx(text, filename):
    """Create DOCX from text"""
    buffer = BytesIO()
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Extracted Text: {filename.rsplit(".", 1)[0]}', 0)
    title.alignment = 1  # Center alignment
    
    # Add some space
    doc.add_paragraph()
    
    # Add text content
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@text_ocr_bp.route('/text-ocr', methods=['GET', 'POST'])
def text_ocr():
    from config import ALLOWED_CONVERTER_EXTENSIONS
    
    # Allowed image formats for OCR
    allowed_extensions = ALLOWED_CONVERTER_EXTENSIONS['image']
    
    if request.method == 'POST':
        # Check if processing text input or file upload
        processing_type = request.form.get('processing_type', 'ocr')
        
        if processing_type == 'text_convert':
            # Handle text conversion
            text_input = request.form.get('text_input', '').strip()
            if not text_input:
                flash('Please enter some text to convert.', 'error')
                return redirect(request.url)
            
            output_format = request.form.get('text_format')
            try:
                if output_format == 'pdf':
                    pdf_buffer = create_text_pdf(text_input, 'text_document')
                    flash('Successfully converted text to PDF!', 'success')
                    return send_file(pdf_buffer, as_attachment=True, download_name="converted_text.pdf")
                elif output_format == 'docx':
                    docx_buffer = create_text_docx(text_input, 'text_document')
                    flash('Successfully converted text to DOCX!', 'success')
                    return send_file(docx_buffer, as_attachment=True, download_name="converted_text.docx")
                elif output_format == 'txt':
                    txt_buffer = BytesIO(text_input.encode('utf-8'))
                    flash('Successfully converted to TXT!', 'success')
                    return send_file(txt_buffer, as_attachment=True, download_name="converted_text.txt")
                elif output_format == 'json':
                    json_data = {
                        "content": text_input,
                        "word_count": len(text_input.split()),
                        "character_count": len(text_input),
                        "line_count": len(text_input.split('\n'))
                    }
                    json_buffer = BytesIO(json.dumps(json_data, indent=2).encode('utf-8'))
                    flash('Successfully converted to JSON!', 'success')
                    return send_file(json_buffer, as_attachment=True, download_name="text_data.json")
            except Exception as e:
                flash(f"Error converting text: {e}", 'error')
                return redirect(request.url)
        
        elif processing_type == 'ocr':
            # Handle OCR processing
            if 'file' not in request.files:
                flash('No file part', 'error')
                return redirect(request.url)
            file = request.files['file']
            if file.filename == '':
                flash('No selected file', 'error')
                return redirect(request.url)

            if file and allowed_file(file.filename, allowed_extensions):
                filename = secure_filename(file.filename)
                filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                
                output_format = request.form.get('ocr_format')
                try:
                    # Extract text using OCR
                    extracted_text = extract_text_from_image(filepath)
                    
                    if not extracted_text:
                        flash('No text found in the image. Please try with a clearer image.', 'warning')
                        return redirect(request.url)
                    
                    if output_format == 'txt':
                        txt_buffer = BytesIO(extracted_text.encode('utf-8'))
                        flash('Successfully extracted text from image!', 'success')
                        return send_file(txt_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}_extracted.txt")
                    elif output_format == 'pdf':
                        pdf_buffer = create_text_pdf(extracted_text, filename)
                        flash('Successfully converted extracted text to PDF!', 'success')
                        return send_file(pdf_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}_extracted.pdf")
                    elif output_format == 'docx':
                        docx_buffer = create_text_docx(extracted_text, filename)
                        flash('Successfully converted extracted text to DOCX!', 'success')
                        return send_file(docx_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}_extracted.docx")
                    elif output_format == 'json':
                        json_data = {
                            "source_image": filename,
                            "extracted_text": extracted_text,
                            "word_count": len(extracted_text.split()),
                            "character_count": len(extracted_text),
                            "line_count": len(extracted_text.split('\n'))
                        }
                        json_buffer = BytesIO(json.dumps(json_data, indent=2).encode('utf-8'))
                        flash('Successfully extracted text to JSON!', 'success')
                        return send_file(json_buffer, as_attachment=True, download_name=f"{filename.rsplit('.', 1)[0]}_extracted.json")
                        
                except Exception as e:
                    flash(f"Error processing image: {e}", 'error')
                    return redirect(request.url)
                finally:
                    # Clean up uploaded file
                    if os.path.exists(filepath):
                        os.remove(filepath)
            else:
                flash('Invalid file type. Please upload an image file (PNG, JPG, JPEG, etc.).', 'error')
                return redirect(request.url)
            
    return render_template('text_ocr.html')

@text_ocr_bp.route('/text-preview', methods=['POST'])
def text_preview():
    """Preview extracted text via AJAX"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    from config import ALLOWED_CONVERTER_EXTENSIONS
    allowed_extensions = ALLOWED_CONVERTER_EXTENSIONS['image']
    
    if file and allowed_file(file.filename, allowed_extensions):
        try:
            # Save file temporarily
            filename = secure_filename(file.filename)
            filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Extract text
            extracted_text = extract_text_from_image(filepath)
            
            # Clean up
            os.remove(filepath)
            
            return jsonify({
                'success': True,
                'text': extracted_text,
                'word_count': len(extracted_text.split()),
                'character_count': len(extracted_text)
            })
            
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    else:
        return jsonify({'error': 'Invalid file type'}), 400
